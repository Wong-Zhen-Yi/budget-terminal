from __future__ import annotations
import base64
import html
import shutil
import sys
from pathlib import Path
from typing import Any

if __package__ in {None, ''}:
    package_root = Path(__file__).resolve().parent.parent
    if str(package_root) not in sys.path:
        sys.path.insert(0, str(package_root))
    from budget_terminal_app import __version__ as APP_VERSION
    from budget_terminal_app.dependencies import *
    from budget_terminal_app.paths import legacy_documents_user_data_path, user_data_path
else:
    from . import __version__ as APP_VERSION
    from .dependencies import *
    from .paths import legacy_documents_user_data_path, user_data_path

DEFAULT_CHART_SLOTS = ['AAPL', 'TSLA', 'NVDA']
USER_DATA_BACKUP_VERSION = 7
BACKUP_BUNDLE_VERSION = 1
USER_DATA_FILE = user_data_path('user_data.json')
LEGACY_USER_DATA_FILE = legacy_documents_user_data_path('user_data.json')
BACKUP_BUNDLE_MANIFEST_NAME = 'manifest.json'
BACKUP_BUNDLE_USER_DATA_NAME = 'user_data.json'
BACKUP_BUNDLE_NOTES_JSON_NAME = 'notes.json'
BACKUP_BUNDLE_NOTES_DOCX_NAME = 'notes.docx'
ROLLBACK_BUNDLES_DIR = user_data_path('backups', 'rollbacks')
DEFAULT_CHART_PAGE_SETTINGS = {
    'symbol': 'SPY',
    'timeframe_label': '1 Day',
    'compare_interval_label': '1 Day',
    'compare_range_label': '5Y',
    'watchlist': [],
    'compare_symbols': [],
    'compare_presets': [],
    'multi_interval_labels': [],
    'indicators': ['Volume', '200 MA'],
    'auto': True,
}
DEFAULT_FUNDAMENTALS_PAGE_SETTINGS = {
    'last_ticker': '',
    'selected_configuration': 'default',
    'custom_selections_by_ticker': {},
}
DEFAULT_DASHBOARD_CHART_SETTINGS = {'symbol': 'SPY', 'timeframe_label': '1 Day', 'indicators': ['Volume', '200 MA'], 'auto': True, 'splitter_sizes': [5, 2], 'main_splitter_sizes': [3, 5]}
DEFAULT_STOCKS_PAGE_SETTINGS = {
    'symbol': 'SPY',
    'auto': True,
    'mfi_enabled': False,
    'main_splitter_sizes': [3, 3, 5],
    'left_splitter_sizes': [4, 2, 3],
    'middle_splitter_sizes': [2, 3],
}
DEFAULT_PORTFOLIO_METRICS_SETTINGS = {'benchmark_symbol': 'SPY', 'lookback_key': '1y'}
DEFAULT_MULTI_CHARTS_SETTINGS = {'custom_symbols': [], 'order': []}
DEFAULT_YOUTUBE_SETTINGS = {'sort_column': -1, 'sort_descending': False}
DEFAULT_THEME_SETTINGS = {'selected_theme': 'trading_dark'}
DEFAULT_OPTIONS_CHAIN_SETTINGS = {'default_risk_free_rate': 0.04}
DEFAULT_NOTES = []
NOTE_CATEGORIES = ('General', 'Observations', 'Trade Ideas')
NOTES_BACKUP_VERSION = 1
MAX_PORTFOLIOS = 5
MULTI_PORTFOLIO_VERSION = 3
PORTFOLIO_IDS = [f'portfolio_{index}' for index in range(1, MAX_PORTFOLIOS + 1)]
DEFAULT_MAIN_PORTFOLIO_ID = PORTFOLIO_IDS[0]
DEFAULT_PORTFOLIO_NAMES = {portfolio_id: f'Portfolio {index}' for index, portfolio_id in enumerate(PORTFOLIO_IDS, start=1)}
SUPPORTED_THEME_IDS = (DEFAULT_THEME_SETTINGS['selected_theme'],)
PORTFOLIO_METRICS_LOOKBACK_CHOICES = ('1y', '3y', '5y', 'max')


def _read_json(path: Any, default: Any) -> Any:
    """Read JSON from disk, returning a fallback on failure."""
    try:
        with Path(path).open() as f:
            return json.load(f)
    except Exception:
        return default


def _write_json(path: Any, data: Any, *, indent: Any=None) -> None:
    """Write JSON data to disk."""
    target = Path(path)
    target.parent.mkdir(parents=True, exist_ok=True)
    temp_path = target.with_suffix(f'{target.suffix}.tmp')
    with temp_path.open('w', encoding='utf-8') as f:
        json.dump(data, f, indent=indent)
    temp_path.replace(target)


def _count_note_images(notes_payload: Any) -> int:
    """Count embedded or referenced note images in a notes payload."""
    total = 0
    if not isinstance(notes_payload, list):
        return 0
    for note in notes_payload:
        if not isinstance(note, dict):
            continue
        images = note.get('images', [])
        if isinstance(images, list):
            total += sum(1 for image in images if isinstance(image, dict))
    return total


def _portfolio_payload_with_chart_slots(data: Any, chart_slots: Any=None) -> Any:
    """Normalize saved portfolio payload into the current on-disk shape."""
    slots = list(chart_slots) if chart_slots else list(DEFAULT_CHART_SLOTS)
    if isinstance(data, dict):
        portfolio = data.get('portfolio', [])
        saved_slots = data.get('chart_slots')
        return {'portfolio': portfolio, 'chart_slots': list(saved_slots) if saved_slots else slots}
    if isinstance(data, list):
        return {'portfolio': data, 'chart_slots': slots}
    return {'portfolio': [], 'chart_slots': slots}


def _normalize_portfolio_id(value: Any) -> Any:
    """Normalize a requested portfolio id into a known fixed slot id."""
    text = str(value or '').strip()
    return text if text in PORTFOLIO_IDS else DEFAULT_MAIN_PORTFOLIO_ID


def _normalize_portfolio_order(raw_order: Any, raw_portfolios: Any=None) -> Any:
    """Normalize portfolio order into a non-empty ordered subset of supported ids."""
    order = []
    if isinstance(raw_order, list):
        for value in raw_order:
            portfolio_id = _normalize_portfolio_id(value)
            if portfolio_id not in order:
                order.append(portfolio_id)
    portfolios = raw_portfolios if isinstance(raw_portfolios, dict) else {}
    for portfolio_id in portfolios.keys():
        clean_id = _normalize_portfolio_id(portfolio_id)
        if clean_id not in order:
            order.append(clean_id)
    if not order:
        order = [DEFAULT_MAIN_PORTFOLIO_ID]
    return order[:MAX_PORTFOLIOS]


def _normalize_selected_portfolio_id(value: Any, portfolio_order: Any, fallback: Any=None) -> Any:
    """Return a selected portfolio id that exists in the ordered portfolio catalog."""
    order = _normalize_portfolio_order(portfolio_order)
    fallback_id = _normalize_portfolio_id(fallback or (order[0] if order else DEFAULT_MAIN_PORTFOLIO_ID))
    portfolio_id = _normalize_portfolio_id(value)
    return portfolio_id if portfolio_id in order else fallback_id


def _normalize_theme_setting(value: Any) -> Any:
    """Clamp persisted theme ids to the currently supported user-selectable set."""
    text = str(value or '').strip()
    return text if text in SUPPORTED_THEME_IDS else DEFAULT_THEME_SETTINGS['selected_theme']


def _normalize_unique_symbol_list(values: Any) -> Any:
    """Normalize ticker-like values into an uppercase unique list."""
    normalized = []
    if not isinstance(values, list):
        return normalized
    for value in values:
        text = str(value or '').upper().strip()
        if text and text not in normalized:
            normalized.append(text)
    return normalized


def _normalize_compare_preset_name(value: Any) -> str:
    """Normalize one compare preset name without losing display casing."""
    return str(value or '').strip()


def _normalize_compare_presets(values: Any) -> list[dict[str, Any]]:
    """Normalize named compare presets into a stable persisted shape."""
    normalized = []
    seen = set()
    if not isinstance(values, list):
        return normalized
    for entry in values:
        preset = entry if isinstance(entry, dict) else {}
        name = _normalize_compare_preset_name(preset.get('name'))
        if not name:
            continue
        name_key = name.casefold()
        if name_key in seen:
            continue
        symbols = _normalize_unique_symbol_list(preset.get('symbols', []))
        if not symbols:
            continue
        interval_label = str(
            preset.get('interval_label', DEFAULT_CHART_PAGE_SETTINGS['compare_interval_label'])
            or DEFAULT_CHART_PAGE_SETTINGS['compare_interval_label']
        ).strip()
        if interval_label not in {'1 Day', '1 Week'}:
            interval_label = DEFAULT_CHART_PAGE_SETTINGS['compare_interval_label']
        range_label = str(
            preset.get('range_label', DEFAULT_CHART_PAGE_SETTINGS['compare_range_label'])
            or DEFAULT_CHART_PAGE_SETTINGS['compare_range_label']
        ).strip().upper()
        if range_label not in {'5Y', '3Y', '1Y', 'YTD', '3M', '1M'}:
            range_label = DEFAULT_CHART_PAGE_SETTINGS['compare_range_label']
        normalized.append({
            'name': name,
            'symbols': symbols,
            'interval_label': interval_label,
            'range_label': range_label,
        })
        seen.add(name_key)
    return normalized


def _normalize_chart_slots_impl(chart_slots: Any, *, allow_empty: bool) -> Any:
    """Normalize chart slots into the persisted three-slot shape."""
    slots = _normalize_unique_symbol_list(chart_slots)
    if not slots and not allow_empty:
        slots = list(DEFAULT_CHART_SLOTS)
    while len(slots) < len(DEFAULT_CHART_SLOTS):
        slots.append(DEFAULT_CHART_SLOTS[len(slots)])
    return slots[:len(DEFAULT_CHART_SLOTS)]


def _normalize_chart_slots(chart_slots: Any) -> Any:
    """Normalize chart slots into the persisted three-slot shape."""
    return _normalize_chart_slots_impl(chart_slots, allow_empty=False)


def _normalize_portfolio_names(raw_names: Any) -> Any:
    """Normalize saved portfolio names for all fixed portfolio slots."""
    names = {}
    raw = raw_names if isinstance(raw_names, dict) else {}
    for index, portfolio_id in enumerate(PORTFOLIO_IDS):
        fallback = f'Portfolio {index + 1}'
        text = str(raw.get(portfolio_id, fallback) or fallback).strip()
        names[portfolio_id] = text or fallback
    return names


def _empty_multi_portfolio_payload(chart_slots: Any=None) -> Any:
    """Create an empty normalized multi-portfolio payload."""
    return {'chart_slots': _normalize_chart_slots(chart_slots), 'portfolios': {portfolio_id: [] for portfolio_id in PORTFOLIO_IDS}}


def _normalize_multi_ticker_payload(data: Any, chart_slots: Any=None) -> Any:
    """Normalize portfolio tickers into the multi-portfolio schema."""
    payload = _empty_multi_portfolio_payload(chart_slots)
    if isinstance(data, dict) and isinstance(data.get('portfolios'), dict):
        payload['chart_slots'] = _normalize_chart_slots(data.get('chart_slots'))
        for portfolio_id in PORTFOLIO_IDS:
            raw_tickers = data.get('portfolios', {}).get(portfolio_id, [])
            if isinstance(raw_tickers, dict):
                raw_tickers = raw_tickers.get('tickers', [])
            payload['portfolios'][portfolio_id] = _normalize_unique_symbol_list(raw_tickers)
        return payload
    legacy = _portfolio_payload_with_chart_slots(data, chart_slots)
    payload['chart_slots'] = _normalize_chart_slots(legacy.get('chart_slots'))
    payload['portfolios'][DEFAULT_MAIN_PORTFOLIO_ID] = _normalize_unique_symbol_list(legacy.get('portfolio', []))
    return payload


def _normalize_multi_tracker_payload(data: Any) -> Any:
    """Normalize tracker data into the multi-portfolio schema."""
    payload = {'portfolios': {portfolio_id: {} for portfolio_id in PORTFOLIO_IDS}}
    if isinstance(data, dict) and isinstance(data.get('portfolios'), dict):
        for portfolio_id in PORTFOLIO_IDS:
            raw_tracker = data.get('portfolios', {}).get(portfolio_id, {})
            payload['portfolios'][portfolio_id] = dict(raw_tracker) if isinstance(raw_tracker, dict) else {}
        return payload
    payload['portfolios'][DEFAULT_MAIN_PORTFOLIO_ID] = dict(data) if isinstance(data, dict) else {}
    return payload


def _normalize_multi_options_payload(data: Any) -> Any:
    """Normalize options positions into the multi-portfolio schema."""
    payload = {'portfolios': {portfolio_id: [] for portfolio_id in PORTFOLIO_IDS}}
    if isinstance(data, dict) and isinstance(data.get('portfolios'), dict):
        for portfolio_id in PORTFOLIO_IDS:
            raw_options = data.get('portfolios', {}).get(portfolio_id, [])
            payload['portfolios'][portfolio_id] = list(raw_options) if isinstance(raw_options, list) else []
        return payload
    payload['portfolios'][DEFAULT_MAIN_PORTFOLIO_ID] = list(data) if isinstance(data, list) else []
    return payload


def _normalize_portfolio_state(state: Any, chart_slots: Any=None) -> Any:
    """Normalize runtime portfolio state into the canonical three-portfolio structure."""
    base = _empty_multi_portfolio_payload(chart_slots)
    raw_state = state if isinstance(state, dict) else {}
    manager = raw_state.get('portfolio_manager', raw_state)
    names = _normalize_portfolio_names(raw_state.get('names', manager.get('names')))
    main_portfolio_id = _normalize_portfolio_id(raw_state.get('main_portfolio_id', manager.get('main_portfolio_id')))
    active_portfolio_id = _normalize_portfolio_id(raw_state.get('active_portfolio_id', manager.get('active_portfolio_id', main_portfolio_id)))
    base['chart_slots'] = _normalize_chart_slots(raw_state.get('chart_slots', chart_slots))
    raw_portfolios = raw_state.get('portfolios', {})
    if not isinstance(raw_portfolios, dict):
        raw_portfolios = {}
    portfolios = {}
    for portfolio_id in PORTFOLIO_IDS:
        raw_entry = raw_portfolios.get(portfolio_id, {})
        if not isinstance(raw_entry, dict):
            raw_entry = {}
        tracker_data = dict(raw_entry.get('tracker_data', {})) if isinstance(raw_entry.get('tracker_data', {}), dict) else {}
        options_data = list(raw_entry.get('options_data', [])) if isinstance(raw_entry.get('options_data', []), list) else []
        portfolios[portfolio_id] = {
            'name': names[portfolio_id],
            'tickers': _normalize_unique_symbol_list(raw_entry.get('tickers', raw_entry.get('portfolio', []))),
            'tracker_data': tracker_data,
            'options_data': options_data,
        }
    return {
        'chart_slots': base['chart_slots'],
        'main_portfolio_id': main_portfolio_id,
        'active_portfolio_id': active_portfolio_id,
        'portfolios': portfolios,
    }


def _sanitize_chart_slots(chart_slots: Any=None) -> Any:
    """Normalize chart-slot input into the persisted 3-slot shape."""
    return _normalize_chart_slots_impl(chart_slots, allow_empty=True)


def _default_portfolio_entry(portfolio_id: Any, chart_slots: Any=None) -> Any:
    """Build an empty normalized portfolio entry."""
    return {
        'id': portfolio_id,
        'name': DEFAULT_PORTFOLIO_NAMES.get(portfolio_id, str(portfolio_id or DEFAULT_MAIN_PORTFOLIO_ID)),
        'portfolio': [],
        'chart_slots': _sanitize_chart_slots(chart_slots),
        'portfolio_tracker': {},
        'options_tracker': [],
    }


def _normalize_portfolio_name(portfolio_id: Any, name: Any) -> Any:
    """Return a safe user-visible portfolio name."""
    text = str(name or '').strip()
    return text if text else DEFAULT_PORTFOLIO_NAMES.get(portfolio_id, str(portfolio_id or DEFAULT_MAIN_PORTFOLIO_ID))


def _normalize_portfolio_catalog(raw_catalog: Any, portfolio_order: Any=None) -> Any:
    """Normalize metadata for the ordered set of existing portfolios."""
    catalog = {}
    source = raw_catalog if isinstance(raw_catalog, dict) else {}
    order = _normalize_portfolio_order(portfolio_order, source)
    for portfolio_id in order:
        entry = source.get(portfolio_id, {})
        if not isinstance(entry, dict):
            entry = {}
        catalog[portfolio_id] = {
            'id': portfolio_id,
            'name': _normalize_portfolio_name(portfolio_id, entry.get('name')),
        }
    return catalog


def _normalize_main_portfolio_id(portfolio_id: Any) -> Any:
    """Clamp a selected portfolio id to a supported fixed slot."""
    portfolio_id = str(portfolio_id or '').strip()
    return portfolio_id if portfolio_id in PORTFOLIO_IDS else DEFAULT_MAIN_PORTFOLIO_ID


def _normalize_portfolio_storage(data: Any) -> Any:
    """Normalize portfolio.json data into the multi-portfolio schema."""
    normalized = {portfolio_id: _default_portfolio_entry(portfolio_id) for portfolio_id in PORTFOLIO_IDS}
    main_portfolio_id = DEFAULT_MAIN_PORTFOLIO_ID
    if isinstance(data, dict) and isinstance(data.get('portfolios'), dict):
        main_portfolio_id = _normalize_main_portfolio_id(data.get('main_portfolio_id'))
        for portfolio_id in PORTFOLIO_IDS:
            entry = data['portfolios'].get(portfolio_id, {})
            if not isinstance(entry, dict):
                entry = {}
            payload = _portfolio_payload_with_chart_slots(entry, normalized[portfolio_id]['chart_slots'])
            normalized[portfolio_id]['portfolio'] = list(payload.get('portfolio', []))
            normalized[portfolio_id]['chart_slots'] = _sanitize_chart_slots(payload.get('chart_slots'))
    elif isinstance(data, dict) or isinstance(data, list):
        payload = _portfolio_payload_with_chart_slots(data)
        normalized[DEFAULT_MAIN_PORTFOLIO_ID]['portfolio'] = list(payload.get('portfolio', []))
        normalized[DEFAULT_MAIN_PORTFOLIO_ID]['chart_slots'] = _sanitize_chart_slots(payload.get('chart_slots'))
    return {'version': MULTI_PORTFOLIO_VERSION, 'main_portfolio_id': main_portfolio_id, 'portfolios': normalized}


def _normalize_tracker_storage(data: Any) -> Any:
    """Normalize portfolio_tracker.json data into the multi-portfolio schema."""
    normalized = {portfolio_id: {} for portfolio_id in PORTFOLIO_IDS}
    if isinstance(data, dict) and isinstance(data.get('portfolios'), dict):
        for portfolio_id in PORTFOLIO_IDS:
            tracker = data['portfolios'].get(portfolio_id, {})
            normalized[portfolio_id] = tracker if isinstance(tracker, dict) else {}
    elif isinstance(data, dict):
        normalized[DEFAULT_MAIN_PORTFOLIO_ID] = data
    return {'version': MULTI_PORTFOLIO_VERSION, 'portfolios': normalized}


def _normalize_options_storage(data: Any) -> Any:
    """Normalize options_tracker.json data into the multi-portfolio schema."""
    normalized = {portfolio_id: [] for portfolio_id in PORTFOLIO_IDS}
    if isinstance(data, dict) and isinstance(data.get('portfolios'), dict):
        for portfolio_id in PORTFOLIO_IDS:
            options_payload = data['portfolios'].get(portfolio_id, [])
            normalized[portfolio_id] = list(options_payload) if isinstance(options_payload, list) else []
    elif isinstance(data, list):
        normalized[DEFAULT_MAIN_PORTFOLIO_ID] = list(data)
    return {'version': MULTI_PORTFOLIO_VERSION, 'portfolios': normalized}


def _normalize_multi_portfolio_state(payload: Any, chart_slots: Any=None) -> Any:
    """Normalize a mixed old/new payload into the runtime multi-portfolio shape."""
    normalized = {
        'version': MULTI_PORTFOLIO_VERSION,
        'main_portfolio_id': DEFAULT_MAIN_PORTFOLIO_ID,
        'active_portfolio_id': DEFAULT_MAIN_PORTFOLIO_ID,
        'portfolio_order': [DEFAULT_MAIN_PORTFOLIO_ID],
        'portfolios': {DEFAULT_MAIN_PORTFOLIO_ID: _default_portfolio_entry(DEFAULT_MAIN_PORTFOLIO_ID, chart_slots)},
    }
    if not isinstance(payload, dict):
        return normalized
    portfolio_map = payload.get('portfolios', {})
    if not isinstance(portfolio_map, dict):
        portfolio_map = {}
    order = _normalize_portfolio_order(payload.get('portfolio_order'), payload.get('portfolio_catalog') or portfolio_map)
    metadata = _normalize_portfolio_catalog(payload.get('portfolio_catalog') or portfolio_map, order)
    normalized['portfolio_order'] = order
    normalized['main_portfolio_id'] = _normalize_selected_portfolio_id(payload.get('main_portfolio_id'), order, order[0])
    normalized['active_portfolio_id'] = _normalize_selected_portfolio_id(payload.get('active_portfolio_id', normalized['main_portfolio_id']), order, normalized['main_portfolio_id'])
    normalized['portfolios'] = {}
    for portfolio_id in order:
        meta_entry = metadata.get(portfolio_id, {})
        raw_entry = portfolio_map.get(portfolio_id, {})
        if not isinstance(raw_entry, dict):
            raw_entry = {}
        portfolio_payload = _portfolio_payload_with_chart_slots(raw_entry.get('portfolio', raw_entry), raw_entry.get('chart_slots'))
        tracker_payload = raw_entry.get('portfolio_tracker', raw_entry.get('tracker_data', {}))
        options_payload = raw_entry.get('options_tracker', raw_entry.get('options_data', []))
        normalized['portfolios'][portfolio_id] = {
            'id': portfolio_id,
            'name': _normalize_portfolio_name(portfolio_id, meta_entry.get('name') or raw_entry.get('name')),
            'portfolio': list(portfolio_payload.get('portfolio', [])),
            'chart_slots': _sanitize_chart_slots(portfolio_payload.get('chart_slots')),
            'portfolio_tracker': tracker_payload if isinstance(tracker_payload, dict) else {},
            'options_tracker': list(options_payload) if isinstance(options_payload, list) else [],
        }
    if 'portfolio' in payload or 'portfolio_tracker' in payload or 'options_tracker' in payload:
        legacy_portfolio = _portfolio_payload_with_chart_slots(payload.get('portfolio', []), chart_slots)
        legacy_tracker = payload.get('portfolio_tracker', {})
        legacy_options = payload.get('options_tracker', [])
        if DEFAULT_MAIN_PORTFOLIO_ID not in normalized['portfolio_order']:
            normalized['portfolio_order'].insert(0, DEFAULT_MAIN_PORTFOLIO_ID)
        normalized['portfolios'][DEFAULT_MAIN_PORTFOLIO_ID] = {
            'id': DEFAULT_MAIN_PORTFOLIO_ID,
            'name': normalized['portfolios'].get(DEFAULT_MAIN_PORTFOLIO_ID, {}).get('name', _normalize_portfolio_name(DEFAULT_MAIN_PORTFOLIO_ID, None)),
            'portfolio': list(legacy_portfolio.get('portfolio', [])),
            'chart_slots': _sanitize_chart_slots(legacy_portfolio.get('chart_slots')),
            'portfolio_tracker': legacy_tracker if isinstance(legacy_tracker, dict) else {},
            'options_tracker': list(legacy_options) if isinstance(legacy_options, list) else [],
        }
        normalized['main_portfolio_id'] = _normalize_selected_portfolio_id(payload.get('main_portfolio_id'), normalized['portfolio_order'], DEFAULT_MAIN_PORTFOLIO_ID)
        normalized['active_portfolio_id'] = _normalize_selected_portfolio_id(payload.get('active_portfolio_id', normalized['main_portfolio_id']), normalized['portfolio_order'], normalized['main_portfolio_id'])
    return normalized


def _serialize_portfolio_storage(state: Any) -> Any:
    """Build the on-disk portfolio.json payload from normalized state."""
    normalized = _normalize_multi_portfolio_state(state)
    return {
        'version': MULTI_PORTFOLIO_VERSION,
        'main_portfolio_id': normalized['main_portfolio_id'],
        'active_portfolio_id': normalized['active_portfolio_id'],
        'portfolio_order': list(normalized.get('portfolio_order', [DEFAULT_MAIN_PORTFOLIO_ID])),
        'portfolios': {
            portfolio_id: _portfolio_payload_with_chart_slots(
                normalized.get('portfolios', {}).get(portfolio_id, {}),
                normalized.get('portfolios', {}).get(portfolio_id, {}).get('chart_slots'),
            ) for portfolio_id in normalized.get('portfolio_order', [DEFAULT_MAIN_PORTFOLIO_ID])
        },
    }


def _serialize_tracker_storage(state: Any) -> Any:
    """Build the on-disk portfolio_tracker.json payload from normalized state."""
    normalized = _normalize_multi_portfolio_state(state)
    return {
        'version': MULTI_PORTFOLIO_VERSION,
        'portfolio_order': list(normalized.get('portfolio_order', [DEFAULT_MAIN_PORTFOLIO_ID])),
        'portfolios': {
            portfolio_id: dict(normalized.get('portfolios', {}).get(portfolio_id, {}).get('portfolio_tracker', {})) for portfolio_id in normalized.get('portfolio_order', [DEFAULT_MAIN_PORTFOLIO_ID])
        },
    }


def _serialize_options_storage(state: Any) -> Any:
    """Build the on-disk options_tracker.json payload from normalized state."""
    normalized = _normalize_multi_portfolio_state(state)
    return {
        'version': MULTI_PORTFOLIO_VERSION,
        'portfolio_order': list(normalized.get('portfolio_order', [DEFAULT_MAIN_PORTFOLIO_ID])),
        'portfolios': {
            portfolio_id: list(normalized.get('portfolios', {}).get(portfolio_id, {}).get('options_tracker', [])) for portfolio_id in normalized.get('portfolio_order', [DEFAULT_MAIN_PORTFOLIO_ID])
        },
    }


def _normalize_networth_payload(data: Any) -> Any:
    """Normalize persisted net-worth lists into the supported shape."""
    payload = data if isinstance(data, dict) else {}
    cash = payload.get('cash', [])
    pension_insurance = payload.get('pension_insurance', [])
    debt = payload.get('debt', [])
    return {
        'cash': list(cash) if isinstance(cash, list) else [],
        'pension_insurance': list(pension_insurance) if isinstance(pension_insurance, list) else [],
        'debt': list(debt) if isinstance(debt, list) else [],
    }


def _normalize_theme_payload(settings: Any) -> Any:
    """Normalize persisted theme settings for the single-file document."""
    saved = settings if isinstance(settings, dict) else {}
    return {'selected_theme': _normalize_theme_setting(saved.get('selected_theme'))}


def _normalize_options_chain_payload(settings: Any) -> Any:
    """Normalize persisted options-chain defaults for the single-file document."""
    saved = settings if isinstance(settings, dict) else {}
    rate = saved.get('default_risk_free_rate', DEFAULT_OPTIONS_CHAIN_SETTINGS['default_risk_free_rate'])
    try:
        rate_value = float(rate)
    except (TypeError, ValueError):
        rate_value = DEFAULT_OPTIONS_CHAIN_SETTINGS['default_risk_free_rate']
    return {'default_risk_free_rate': min(max(rate_value, 0.0), 1.0)}


def _normalize_note_category(value: Any) -> str:
    """Clamp note categories to the supported fixed set."""
    text = str(value or '').strip()
    for category in NOTE_CATEGORIES:
        if text.casefold() == category.casefold():
            return category
    return NOTE_CATEGORIES[0]


def _normalize_note_images(images: Any) -> list[dict[str, str]]:
    """Normalize persisted note attachments into a compact list."""
    normalized = []
    if not isinstance(images, list):
        return normalized
    for image in images:
        entry = image if isinstance(image, dict) else {'path': image}
        path_text = str(entry.get('path', '') or '').strip().replace('\\', '/')
        if not path_text:
            continue
        image_id = str(entry.get('id', '') or '').strip() or path_text
        name = str(entry.get('name', '') or Path(path_text).name).strip() or Path(path_text).name
        normalized.append({'id': image_id, 'name': name, 'path': path_text})
    return normalized


def _normalize_notes_payload(notes: Any) -> list[dict[str, Any]]:
    """Normalize persisted notes into the canonical list shape."""
    normalized = []
    if not isinstance(notes, list):
        return normalized
    seen_ids = set()
    for index, note in enumerate(notes):
        if not isinstance(note, dict):
            continue
        note_id = str(note.get('id', '') or '').strip() or f'note_{index + 1}'
        if note_id in seen_ids:
            continue
        seen_ids.add(note_id)
        created_at = str(note.get('created_at', '') or '').strip()
        updated_at = str(note.get('updated_at', '') or '').strip() or created_at
        normalized.append({
            'id': note_id,
            'title': str(note.get('title', '') or ''),
            'body': str(note.get('body', '') or ''),
            'category': _normalize_note_category(note.get('category')),
            'created_at': created_at,
            'updated_at': updated_at,
            'images': _normalize_note_images(note.get('images', [])),
        })
    return normalized


def _notes_image_directory() -> Any:
    """Return the on-disk directory that stores copied note images."""
    return user_data_path('notes_images')


def _cleanup_orphaned_note_images(notes: Any) -> None:
    """Remove copied note-image folders that are no longer referenced by saved notes."""
    root = Path(_notes_image_directory())
    if not root.exists():
        return
    valid_note_ids = {str(note.get('id', '') or '').strip() for note in notes if isinstance(note, dict)}
    for child in root.iterdir():
        if not child.is_dir():
            continue
        if child.name not in valid_note_ids:
            shutil.rmtree(child, ignore_errors=True)

def fmt_num(val: Any) -> Any:
    """Format large numbers with B/M/K suffix."""
    if val is None:
        return 'N/A'
    try:
        val = float(val)
        if pd.isna(val):
            return 'N/A'
    except (TypeError, ValueError):
        return 'N/A'
    neg = val < 0
    abs_val = abs(val)
    if abs_val >= 1000000000000.0:
        s = f'{abs_val / 1000000000000.0:.2f}T'
    elif abs_val >= 1000000000.0:
        s = f'{abs_val / 1000000000.0:.2f}B'
    elif abs_val >= 1000000.0:
        s = f'{abs_val / 1000000.0:.2f}M'
    elif abs_val >= 1000.0:
        s = f'{abs_val / 1000.0:.1f}K'
    else:
        s = f'{abs_val:.2f}'
    return f'-{s}' if neg else s


def _default_user_data_document() -> Any:
    """Build the default single-file user-data document."""
    portfolio_state = _normalize_multi_portfolio_state({})
    return {
        'version': USER_DATA_BACKUP_VERSION,
        'main_portfolio_id': portfolio_state['main_portfolio_id'],
        'active_portfolio_id': portfolio_state['active_portfolio_id'],
        'portfolio_order': list(portfolio_state.get('portfolio_order', [DEFAULT_MAIN_PORTFOLIO_ID])),
        'portfolios': portfolio_state['portfolios'],
        'fundamentals_page': _normalize_fundamentals_page_settings(DEFAULT_FUNDAMENTALS_PAGE_SETTINGS),
        'chart_page': DEFAULT_CHART_PAGE_SETTINGS.copy(),
        'dashboard_chart': DEFAULT_DASHBOARD_CHART_SETTINGS.copy(),
        'stocks_page': DEFAULT_STOCKS_PAGE_SETTINGS.copy(),
        'portfolio_metrics': DEFAULT_PORTFOLIO_METRICS_SETTINGS.copy(),
        'multi_charts': DEFAULT_MULTI_CHARTS_SETTINGS.copy(),
        'youtube': DEFAULT_YOUTUBE_SETTINGS.copy(),
        'net_worth': {'cash': [], 'pension_insurance': [], 'debt': []},
        'notes': list(DEFAULT_NOTES),
        'theme': DEFAULT_THEME_SETTINGS.copy(),
        'options_chain': DEFAULT_OPTIONS_CHAIN_SETTINGS.copy(),
        'time_12h': False,
    }


def _normalize_user_data_document(payload: Any, *, existing_notes: Any=None) -> Any:
    """Normalize persisted single-file user data into the canonical shape."""
    default = _default_user_data_document()
    saved = payload if isinstance(payload, dict) else {}
    portfolio_state = _normalize_multi_portfolio_state(saved)
    notes_fallback = existing_notes if existing_notes is not None else default['notes']
    chart_page_payload = saved.get('chart_page', default['chart_page'])
    exported_compare_presets = saved.get('compare_presets')
    if isinstance(exported_compare_presets, list):
        chart_page_payload = dict(chart_page_payload) if isinstance(chart_page_payload, dict) else {}
        chart_page_payload['compare_presets'] = exported_compare_presets
    return {
        'version': USER_DATA_BACKUP_VERSION,
        'main_portfolio_id': portfolio_state['main_portfolio_id'],
        'active_portfolio_id': portfolio_state.get('active_portfolio_id', portfolio_state['main_portfolio_id']),
        'portfolio_order': list(portfolio_state.get('portfolio_order', [DEFAULT_MAIN_PORTFOLIO_ID])),
        'portfolios': portfolio_state['portfolios'],
        'fundamentals_page': _normalize_fundamentals_page_settings(saved.get('fundamentals_page', default['fundamentals_page'])),
        'chart_page': _normalize_chart_page_settings(chart_page_payload),
        'dashboard_chart': _normalize_dashboard_chart_settings(saved.get('dashboard_chart', default['dashboard_chart'])),
        'stocks_page': _normalize_stocks_page_settings(saved.get('stocks_page', default['stocks_page'])),
        'portfolio_metrics': _normalize_portfolio_metrics_settings(saved.get('portfolio_metrics', default['portfolio_metrics'])),
        'multi_charts': _normalize_multi_charts_settings(saved.get('multi_charts', default['multi_charts'])),
        'youtube': _normalize_youtube_settings(saved.get('youtube', default['youtube'])),
        'net_worth': _normalize_networth_payload(saved.get('net_worth', default['net_worth'])),
        'notes': _normalize_notes_payload(saved.get('notes', notes_fallback)),
        'theme': _normalize_theme_payload(saved.get('theme', default['theme'])),
        'options_chain': _normalize_options_chain_payload(saved.get('options_chain', default['options_chain'])),
        'time_12h': bool(saved.get('time_12h', False)),
    }


def _load_legacy_user_data_document() -> Any:
    """Load the old Documents-based user-data file when present."""
    return _normalize_user_data_document(_read_json(LEGACY_USER_DATA_FILE, None))


def _migrate_legacy_user_data_file() -> Any:
    """Move the legacy Documents-based save file into LocalAppData once."""
    if Path(USER_DATA_FILE).exists():
        return None
    legacy_path = Path(LEGACY_USER_DATA_FILE)
    if not legacy_path.exists():
        return None
    normalized = _load_legacy_user_data_document()
    _write_json(USER_DATA_FILE, normalized, indent=2)
    try:
        legacy_path.unlink()
    except OSError:
        logger.warning('User data migrated to %s but the legacy file could not be removed: %s', USER_DATA_FILE, legacy_path)
    else:
        logger.info('Migrated user data from %s to %s.', legacy_path, USER_DATA_FILE)
    return normalized


def _load_user_data_document() -> Any:
    """Load the current single-file user-data document from LocalAppData."""
    migrated = _migrate_legacy_user_data_file()
    if migrated is not None:
        return migrated
    return _normalize_user_data_document(_read_json(USER_DATA_FILE, None))


def _save_user_data_document(data: Any) -> Any:
    """Persist the normalized single-file user-data document to LocalAppData."""
    normalized = _normalize_user_data_document(data)
    _write_json(USER_DATA_FILE, normalized, indent=2)
    _cleanup_orphaned_note_images(normalized.get('notes', []))
    return normalized


def load_all_portfolios_state() -> Any:
    """Load the normalized state for all supported portfolios."""
    document = _load_user_data_document()
    return _normalize_multi_portfolio_state(document)


def save_all_portfolios_state(state: Any) -> Any:
    """Persist the normalized state for all supported portfolios."""
    normalized = _normalize_multi_portfolio_state(state)
    document = _load_user_data_document()
    document['main_portfolio_id'] = normalized['main_portfolio_id']
    document['active_portfolio_id'] = normalized.get('active_portfolio_id', normalized['main_portfolio_id'])
    document['portfolio_order'] = list(normalized.get('portfolio_order', [DEFAULT_MAIN_PORTFOLIO_ID]))
    document['portfolios'] = normalized['portfolios']
    _save_user_data_document(document)
    return normalized


def _resolve_state_portfolio_id(state: Any, portfolio_id: Any=None) -> Any:
    """Return an existing portfolio id from normalized state."""
    order = _normalize_portfolio_order(state.get('portfolio_order'), state.get('portfolios', {})) if isinstance(state, dict) else [DEFAULT_MAIN_PORTFOLIO_ID]
    fallback = state.get('active_portfolio_id') if isinstance(state, dict) else DEFAULT_MAIN_PORTFOLIO_ID
    return _normalize_selected_portfolio_id(portfolio_id or fallback or DEFAULT_MAIN_PORTFOLIO_ID, order, order[0])


def load_active_portfolio_state(portfolio_id: Any=None) -> Any:
    """Load one normalized portfolio state."""
    state = load_all_portfolios_state()
    active_id = _resolve_state_portfolio_id(state, portfolio_id or state.get('active_portfolio_id') or state['main_portfolio_id'])
    active = state['portfolios'][active_id]
    return {
        'main_portfolio_id': state.get('main_portfolio_id', active_id),
        'portfolio_id': active_id,
        'portfolio_name': active.get('name', DEFAULT_PORTFOLIO_NAMES.get(active_id, active_id)),
        'portfolio': list(active.get('portfolio', [])),
        'chart_slots': _sanitize_chart_slots(active.get('chart_slots')),
        'portfolio_tracker': dict(active.get('portfolio_tracker', {})),
        'options_tracker': list(active.get('options_tracker', [])),
    }

def load_tickers() -> Any:
    """Load tickers."""
    state = load_active_portfolio_state()
    tickers = state.get('portfolio', [])
    chart_slots = state.get('chart_slots', list(DEFAULT_CHART_SLOTS))
    if tickers:
        return (tickers, chart_slots)
    return (['AAPL', 'MSFT', 'GOOGL', 'TSLA'], chart_slots)

def save_tickers(portfolio: Any, chart_slots: Any=None, portfolio_id: Any=None) -> None:
    """Save tickers."""
    state = load_all_portfolios_state()
    active_id = _resolve_state_portfolio_id(state, portfolio_id or state.get('active_portfolio_id') or state['main_portfolio_id'])
    state['portfolios'][active_id]['portfolio'] = list(portfolio) if isinstance(portfolio, list) else []
    state['portfolios'][active_id]['chart_slots'] = _sanitize_chart_slots(chart_slots if chart_slots is not None else state['portfolios'][active_id].get('chart_slots'))
    save_all_portfolios_state(state)

def load_tracker_data(portfolio_id: Any=None) -> Any:
    """Load tracker data."""
    return load_active_portfolio_state(portfolio_id).get('portfolio_tracker', {})

def save_tracker_data(data: Any, portfolio_id: Any=None) -> None:
    """Save tracker data."""
    state = load_all_portfolios_state()
    active_id = _resolve_state_portfolio_id(state, portfolio_id or state.get('active_portfolio_id') or state['main_portfolio_id'])
    state['portfolios'][active_id]['portfolio_tracker'] = data if isinstance(data, dict) else {}
    save_all_portfolios_state(state)

def load_options_data(portfolio_id: Any=None) -> Any:
    """Load options data."""
    return load_active_portfolio_state(portfolio_id).get('options_tracker', [])

def save_options_data(data: Any, portfolio_id: Any=None) -> None:
    """Save options data."""
    state = load_all_portfolios_state()
    active_id = _resolve_state_portfolio_id(state, portfolio_id or state.get('active_portfolio_id') or state['main_portfolio_id'])
    state['portfolios'][active_id]['options_tracker'] = list(data) if isinstance(data, list) else []
    save_all_portfolios_state(state)

def load_networth_data() -> Any:
    """Load networth data."""
    return _normalize_networth_payload(_load_user_data_document().get('net_worth'))

def save_networth_data(data: Any) -> None:
    """Save networth data."""
    document = _load_user_data_document()
    document['net_worth'] = _normalize_networth_payload(data)
    _save_user_data_document(document)


def load_notes_data() -> Any:
    """Load persisted notes."""
    return _normalize_notes_payload(_load_user_data_document().get('notes', []))


def save_notes_data(data: Any) -> Any:
    """Persist notes and return the normalized saved list."""
    document = _load_user_data_document()
    document['notes'] = _normalize_notes_payload(data)
    saved = _save_user_data_document(document)
    return list(saved.get('notes', []))


def _validate_notes_backup_payload(payload: Any) -> Any:
    """Validate imported notes-backup data and return the original payload."""
    if not isinstance(payload, dict):
        raise ValueError('Notes backup file must contain a JSON object.')
    if not isinstance(payload.get('notes'), list):
        raise ValueError('Notes backup file must include a notes list.')
    return payload


def build_notes_backup() -> Any:
    """Build a standalone notes backup payload with embedded image bytes."""
    exported_notes = []
    for note in load_notes_data():
        images = []
        for image in _normalize_note_images(note.get('images', [])):
            relative_path = str(image.get('path', '') or '').strip().replace('\\', '/')
            if not relative_path:
                continue
            image_path = Path(user_data_path(*Path(relative_path).parts))
            if not image_path.exists() or not image_path.is_file():
                logger.warning('Skipping missing note image during notes export: %s', image_path)
                continue
            try:
                raw_bytes = image_path.read_bytes()
            except OSError as exc:
                logger.warning('Skipping unreadable note image during notes export %s: %s', image_path, exc)
                continue
            images.append({
                'id': str(image.get('id', '') or '').strip(),
                'name': str(image.get('name', '') or image_path.name).strip() or image_path.name,
                'file_name': image_path.name,
                'data_base64': base64.b64encode(raw_bytes).decode('ascii'),
            })
        exported_notes.append({
            'id': str(note.get('id', '') or '').strip(),
            'title': str(note.get('title', '') or ''),
            'body': str(note.get('body', '') or ''),
            'category': _normalize_note_category(note.get('category')),
            'created_at': str(note.get('created_at', '') or '').strip(),
            'updated_at': str(note.get('updated_at', '') or '').strip(),
            'images': images,
        })
    return {
        'version': NOTES_BACKUP_VERSION,
        'exported_at': datetime.datetime.now(datetime.timezone.utc).isoformat(),
        'notes': exported_notes,
    }


def _build_backup_bundle_manifest(user_data_backup: Any, notes_backup: Any) -> dict[str, Any]:
    """Build metadata for a folder-based backup bundle."""
    portfolio_state = _normalize_multi_portfolio_state(user_data_backup)
    notes = notes_backup.get('notes', []) if isinstance(notes_backup, dict) else []
    return {
        'bundle_version': BACKUP_BUNDLE_VERSION,
        'app_version': APP_VERSION,
        'exported_at': str(user_data_backup.get('exported_at', '') or notes_backup.get('exported_at', '')),
        'files': {
            'user_data': BACKUP_BUNDLE_USER_DATA_NAME,
            'notes': BACKUP_BUNDLE_NOTES_JSON_NAME,
            'notes_docx': BACKUP_BUNDLE_NOTES_DOCX_NAME,
        },
        'counts': {
            'portfolios': len(portfolio_state.get('portfolio_order', [])),
            'notes': len(notes) if isinstance(notes, list) else 0,
            'note_images': _count_note_images(notes),
        },
    }


def export_notes_backup(path: Any) -> None:
    """Write a standalone notes backup to disk."""
    _write_json(path, build_notes_backup(), indent=2)


def load_notes_backup(path: Any) -> Any:
    """Load and validate a standalone notes backup file."""
    source_path = Path(path)
    suffix = source_path.suffix.lower()
    if suffix == '.docx':
        return _load_notes_backup_docx(source_path)
    payload = _read_json(source_path, None)
    if payload is None:
        raise ValueError('Unable to read notes backup file.')
    return _validate_notes_backup_payload(payload)


def _parse_docx_note_meta(meta_text: Any) -> dict[str, str]:
    """Parse a DOCX-exported note metadata block."""
    parsed = {'category': NOTE_CATEGORIES[0], 'created_at': '', 'updated_at': ''}
    for raw_line in str(meta_text or '').splitlines():
        line = str(raw_line or '').strip()
        if not line or ':' not in line:
            continue
        label, value = line.split(':', 1)
        normalized_value = value.strip()
        label_key = label.strip().lower()
        if label_key == 'category':
            parsed['category'] = _normalize_note_category(normalized_value)
        elif label_key == 'created':
            parsed['created_at'] = '' if normalized_value == '-' else normalized_value
        elif label_key == 'edited':
            parsed['updated_at'] = '' if normalized_value == '-' else normalized_value
    if not parsed['updated_at']:
        parsed['updated_at'] = parsed['created_at']
    return parsed


def _docx_paragraph_images(document: Any, paragraph: Any) -> list[dict[str, str]]:
    """Extract embedded image payloads from a DOCX paragraph."""
    image_entries = []
    namespaces = {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}
    rel_attr = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
    for image_index, blip in enumerate(paragraph._p.iterfind('.//a:blip', namespaces), start=1):
        rel_id = str(blip.get(rel_attr, '') or '').strip()
        if not rel_id:
            continue
        image_part = document.part.related_parts.get(rel_id)
        if image_part is None:
            continue
        raw_bytes = getattr(image_part, 'blob', b'') or b''
        if not raw_bytes:
            continue
        part_name = Path(str(getattr(image_part, 'partname', '') or ''))
        file_name = part_name.name or f'image_{image_index}.bin'
        image_entries.append({
            'name': file_name,
            'file_name': file_name,
            'data_base64': base64.b64encode(raw_bytes).decode('ascii'),
        })
    return image_entries


def _finalize_docx_import_note(note: Any, body_lines: Any) -> dict[str, Any]:
    """Convert parsed DOCX note fragments into backup payload shape."""
    body_text = '\n'.join(body_lines) if isinstance(body_lines, list) else str(body_lines or '')
    return {
        'title': str(note.get('title', '') or '') or 'Untitled note',
        'body': body_text,
        'category': _normalize_note_category(note.get('category')),
        'created_at': str(note.get('created_at', '') or '').strip(),
        'updated_at': str(note.get('updated_at', '') or '').strip() or str(note.get('created_at', '') or '').strip(),
        'images': list(note.get('images', [])) if isinstance(note.get('images', []), list) else [],
    }


def _load_notes_backup_docx(path: Any) -> Any:
    """Load a DOCX notes export and convert it into backup payload shape."""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError('DOCX notes import requires python-docx. Install it with: python -m pip install python-docx') from exc

    document = Document(str(path))
    imported_notes = []
    current_note = None
    current_body_lines: list[str] = []
    current_section = None

    for paragraph in document.paragraphs:
        style_name = str(getattr(getattr(paragraph, 'style', None), 'name', '') or '')
        paragraph_text = str(paragraph.text or '')
        stripped_text = paragraph_text.strip()

        if style_name == 'Title':
            continue

        if style_name == 'Heading 1':
            if current_note is not None:
                imported_notes.append(_finalize_docx_import_note(current_note, current_body_lines))
            current_note = {
                'title': stripped_text or 'Untitled note',
                'category': NOTE_CATEGORIES[0],
                'created_at': '',
                'updated_at': '',
                'images': [],
            }
            current_body_lines = []
            current_section = 'meta'
            continue

        if current_note is None:
            continue

        if style_name == 'Heading 2':
            heading_key = stripped_text.casefold()
            if heading_key == 'note':
                current_section = 'body'
            elif heading_key == 'pictures':
                current_section = 'pictures'
            else:
                current_section = None
            continue

        if current_section == 'meta':
            current_note.update(_parse_docx_note_meta(paragraph_text))
            continue

        if current_section == 'body':
            if not current_body_lines and stripped_text == 'No body text.':
                continue
            current_body_lines.append(paragraph_text)
            continue

        if current_section == 'pictures':
            current_note['images'].extend(_docx_paragraph_images(document, paragraph))

    if current_note is not None:
        imported_notes.append(_finalize_docx_import_note(current_note, current_body_lines))

    if not imported_notes:
        non_empty_paragraphs = [str(paragraph.text or '').strip() for paragraph in document.paragraphs if str(paragraph.text or '').strip()]
        if any(text != 'No notes were available at export time.' for text in non_empty_paragraphs):
            raise ValueError('DOCX file does not match the Budget Terminal notes export format.')

    return _validate_notes_backup_payload({
        'version': NOTES_BACKUP_VERSION,
        'imported_at': datetime.datetime.now(datetime.timezone.utc).isoformat(),
        'notes': imported_notes,
    })


def apply_notes_backup(payload: Any) -> Any:
    """Persist a standalone notes backup and return the normalized saved notes list."""
    validated = _validate_notes_backup_payload(payload)
    notes_root = Path(_notes_image_directory())
    notes_root.mkdir(parents=True, exist_ok=True)
    imported_notes = []
    seen_note_ids = set()
    for note_index, raw_note in enumerate(validated.get('notes', [])):
        if not isinstance(raw_note, dict):
            continue
        base_note_id = str(raw_note.get('id', '') or f'note_{note_index + 1}').strip() or f'note_{note_index + 1}'
        note_id = base_note_id
        counter = 2
        while note_id in seen_note_ids:
            note_id = f'{base_note_id}_{counter}'
            counter += 1
        seen_note_ids.add(note_id)
        note_dir = notes_root / note_id
        shutil.rmtree(note_dir, ignore_errors=True)
        images = []
        raw_images = raw_note.get('images', [])
        if isinstance(raw_images, list) and raw_images:
            note_dir.mkdir(parents=True, exist_ok=True)
            for image_index, raw_image in enumerate(raw_images):
                if not isinstance(raw_image, dict):
                    continue
                encoded = str(raw_image.get('data_base64', '') or '').strip()
                if not encoded:
                    continue
                try:
                    raw_bytes = base64.b64decode(encoded, validate=True)
                except Exception:
                    logger.warning('Skipping invalid note image payload for imported note %s.', note_id)
                    continue
                base_name = Path(str(raw_image.get('file_name', '') or raw_image.get('name', '') or f'image_{image_index + 1}.bin')).name
                if not base_name:
                    base_name = f'image_{image_index + 1}.bin'
                target_path = note_dir / base_name
                file_counter = 2
                while target_path.exists():
                    target_path = note_dir / f'{Path(base_name).stem}_{file_counter}{Path(base_name).suffix}'
                    file_counter += 1
                try:
                    target_path.write_bytes(raw_bytes)
                except OSError as exc:
                    logger.warning('Skipping note image write failure for %s: %s', target_path, exc)
                    continue
                image_id = str(raw_image.get('id', '') or f'{note_id}_image_{image_index + 1}').strip() or f'{note_id}_image_{image_index + 1}'
                image_name = str(raw_image.get('name', '') or target_path.name).strip() or target_path.name
                images.append({
                    'id': image_id,
                    'name': image_name,
                    'path': str(Path('notes_images') / note_id / target_path.name).replace('\\', '/'),
                })
        if note_dir.exists() and not any(note_dir.iterdir()):
            shutil.rmtree(note_dir, ignore_errors=True)
        created_at = str(raw_note.get('created_at', '') or '').strip()
        updated_at = str(raw_note.get('updated_at', '') or '').strip() or created_at
        imported_notes.append({
            'id': note_id,
            'title': str(raw_note.get('title', '') or ''),
            'body': str(raw_note.get('body', '') or ''),
            'category': _normalize_note_category(raw_note.get('category')),
            'created_at': created_at,
            'updated_at': updated_at,
            'images': images,
        })
    return save_notes_data(imported_notes)


def _parse_note_timestamp(value: Any) -> Any:
    """Parse a persisted note timestamp into an aware datetime."""
    text = str(value or '').strip()
    if not text:
        return None
    try:
        parsed = datetime.datetime.fromisoformat(text.replace('Z', '+00:00'))
    except ValueError:
        return None
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=datetime.timezone.utc)
    return parsed


def _sorted_notes_for_export() -> list[dict[str, Any]]:
    """Return saved notes ordered newest-first for export outputs."""
    return sorted(
        load_notes_data(),
        key=lambda note: (_parse_note_timestamp(note.get('updated_at')) or _parse_note_timestamp(note.get('created_at')) or datetime.datetime.fromtimestamp(0, datetime.timezone.utc)),
        reverse=True,
    )


def _note_image_path(image: Any) -> Any:
    """Resolve a persisted note image reference into a local path."""
    return Path(user_data_path(*Path(str(image.get('path', '') or '').strip().replace('\\', '/')).parts))


def _note_image_export_entries(note: Any) -> list[dict[str, Any]]:
    """Collect note image export metadata for DOCX/HTML/PDF outputs."""
    entries = []
    for image in _normalize_note_images(note.get('images', [])):
        image_path = _note_image_path(image)
        if not image_path.exists() or not image_path.is_file():
            logger.warning('Skipping missing note image during export: %s', image_path)
            entries.append({'status': 'missing', 'path': image_path})
            continue
        try:
            raw_bytes = image_path.read_bytes()
        except OSError:
            logger.exception('Unable to read note image during export: %s', image_path)
            entries.append({'status': 'missing', 'path': image_path})
            continue
        suffix = image_path.suffix.lower()
        mime_type = {
            '.png': 'image/png',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.gif': 'image/gif',
            '.bmp': 'image/bmp',
            '.webp': 'image/webp',
        }.get(suffix, 'application/octet-stream')
        entries.append({
            'status': 'ok',
            'path': image_path,
            'mime_type': mime_type,
            'data_base64': base64.b64encode(raw_bytes).decode('ascii'),
        })
    return entries


def _build_notes_export_html(*, for_pdf: bool=False) -> str:
    """Render notes as standalone HTML with embedded images."""
    notes = _sorted_notes_for_export()
    timestamp = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    page_break_rule = 'page-break-before: always;' if for_pdf else 'border-top: 1px solid #d7dde5; margin-top: 28px; padding-top: 28px;'
    parts = [
        '<!DOCTYPE html>',
        '<html lang="en">',
        '<head>',
        '<meta charset="utf-8">',
        '<title>Budget Terminal Notes</title>',
        '<style>',
        'body { font-family: Arial, sans-serif; color: #1f2937; margin: 32px; line-height: 1.5; }',
        'h1 { font-size: 24px; margin: 0 0 10px; }',
        'h2 { font-size: 18px; margin: 18px 0 8px; }',
        'h3 { font-size: 14px; margin: 16px 0 6px; text-transform: uppercase; letter-spacing: 0.04em; color: #4b5563; }',
        '.summary { margin-bottom: 24px; color: #4b5563; }',
        '.note { margin-bottom: 24px; }',
        f'.note + .note {{ {page_break_rule} }}',
        '.meta { margin: 10px 0 16px; color: #374151; }',
        '.body-line { margin: 0 0 8px; white-space: pre-wrap; }',
        '.image { margin: 12px 0 16px; }',
        '.image img { max-width: 100%; max-height: 560px; border: 1px solid #d7dde5; }',
        '.placeholder { margin: 12px 0; color: #6b7280; font-style: italic; }',
        '</style>',
        '</head>',
        '<body>',
        '<h1>Budget Terminal Notes</h1>',
        f'<div class="summary"><strong>Exported at:</strong> {html.escape(timestamp)}<br><strong>Note count:</strong> {len(notes)}</div>',
    ]
    if not notes:
        parts.append('<p>No notes were available at export time.</p>')
    for note in notes:
        parts.append('<section class="note">')
        parts.append(f'<h2>{html.escape(str(note.get("title", "") or "Untitled note"))}</h2>')
        parts.append(
            '<div class="meta">'
            f'<strong>Category:</strong> {html.escape(str(note.get("category", NOTE_CATEGORIES[0]) or NOTE_CATEGORIES[0]))}<br>'
            f'<strong>Created:</strong> {html.escape(str(note.get("created_at", "") or "-"))}<br>'
            f'<strong>Edited:</strong> {html.escape(str(note.get("updated_at", "") or "-"))}'
            '</div>'
        )
        parts.append('<h3>Note</h3>')
        body_text = str(note.get('body', '') or '')
        if body_text.strip():
            for line in body_text.splitlines():
                parts.append(f'<div class="body-line">{html.escape(line) if line.strip() else "&nbsp;"}</div>')
        else:
            parts.append('<div class="body-line">No body text.</div>')
        image_entries = _note_image_export_entries(note)
        if image_entries:
            parts.append('<h3>Pictures</h3>')
        for image_entry in image_entries:
            if image_entry.get('status') != 'ok':
                parts.append('<div class="placeholder">Picture could not be embedded.</div>')
                continue
            parts.append(
                '<div class="image">'
                f'<img src="data:{html.escape(str(image_entry.get("mime_type", "image/png")))};base64,{image_entry.get("data_base64", "")}" alt="Embedded note picture">'
                '</div>'
            )
        parts.append('</section>')
    parts.extend(['</body>', '</html>'])
    return '\n'.join(parts)


def export_notes_docx(path: Any) -> None:
    """Write all saved notes to a DOCX document."""
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as exc:
        raise RuntimeError('Notes export requires python-docx. Install it with: python -m pip install python-docx') from exc

    notes = _sorted_notes_for_export()

    document = Document()
    styles = document.styles
    if 'Normal' in styles:
        styles['Normal'].font.name = 'Arial'
        styles['Normal'].font.size = None
    document.core_properties.title = 'Budget Terminal Notes Export'
    document.core_properties.subject = 'Exported notes'
    document.core_properties.comments = 'Generated by Budget Terminal'

    document.add_heading('Budget Terminal Notes', 0)
    summary = document.add_paragraph()
    summary.add_run('Exported at: ').bold = True
    summary.add_run(datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    summary.add_run('\n')
    summary.add_run('Note count: ').bold = True
    summary.add_run(str(len(notes)))

    if not notes:
        document.add_paragraph('No notes were available at export time.')
        document.save(path)
        return

    for index, note in enumerate(notes):
        if index > 0:
            document.add_page_break()
        document.add_heading(str(note.get('title', '') or 'Untitled note'), level=1)

        meta = document.add_paragraph()
        meta.add_run('Category: ').bold = True
        meta.add_run(str(note.get('category', NOTE_CATEGORIES[0]) or NOTE_CATEGORIES[0]))
        meta.add_run('\n')
        meta.add_run('Created: ').bold = True
        meta.add_run(str(note.get('created_at', '') or '-'))
        meta.add_run('\n')
        meta.add_run('Edited: ').bold = True
        meta.add_run(str(note.get('updated_at', '') or '-'))

        body_text = str(note.get('body', '') or '')
        document.add_heading('Note', level=2)
        if body_text.strip():
            for line in body_text.splitlines():
                document.add_paragraph(line if line.strip() else '')
        else:
            document.add_paragraph('No body text.')

        image_entries = _note_image_export_entries(note)
        if image_entries:
            document.add_heading('Pictures', level=2)
        for image_entry in image_entries:
            if image_entry.get('status') != 'ok':
                document.add_paragraph('Picture could not be embedded.')
                continue
            try:
                image_path = Path(image_entry.get('path'))
                picture = document.add_picture(str(image_path), width=Inches(5.8))
                inline = picture._inline
                inline.graphic.graphicData.pic.nvPicPr.cNvPr.set('name', 'Embedded Picture')
            except Exception:
                logger.exception('Unable to embed note image in DOCX export: %s', image_path)
                document.add_paragraph('Picture could not be embedded.')

    document.save(path)


def export_notes_html(path: Any) -> None:
    """Write all saved notes to a standalone HTML document."""
    Path(path).write_text(_build_notes_export_html(for_pdf=False), encoding='utf-8')


def build_user_data_backup(*, include_notes: bool=True) -> Any:
    """Build a single-file backup payload for all persisted user data."""
    backup = _load_user_data_document()
    if not include_notes:
        backup = dict(backup)
        backup.pop('notes', None)
    backup['compare_presets'] = list(
        _normalize_chart_page_settings(backup.get('chart_page', DEFAULT_CHART_PAGE_SETTINGS)).get('compare_presets', [])
    )
    backup['version'] = USER_DATA_BACKUP_VERSION
    backup['exported_at'] = datetime.datetime.now(datetime.timezone.utc).isoformat()
    return backup


def export_user_data_backup(path: Any) -> None:
    """Write the current backup payload to a single JSON file."""
    _write_json(path, build_user_data_backup(), indent=2)


def _bundle_directory(parent_directory: Any, *, prefix: str='budget_terminal_backup') -> Path:
    """Create a unique bundle directory within the requested parent."""
    target_root = Path(parent_directory)
    target_root.mkdir(parents=True, exist_ok=True)
    base_name = f"{prefix}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}"
    bundle_dir = target_root / base_name
    suffix = 2
    while bundle_dir.exists():
        bundle_dir = target_root / f'{base_name}_{suffix}'
        suffix += 1
    bundle_dir.mkdir(parents=True, exist_ok=False)
    return bundle_dir


def export_user_data_bundle(parent_directory: Any, *, prefix: str='budget_terminal_backup') -> dict[str, str]:
    """Write separate user-data JSON and notes DOCX exports into a timestamped folder."""
    bundle_dir = _bundle_directory(parent_directory, prefix=prefix)
    user_data_backup = build_user_data_backup(include_notes=False)
    notes_backup = build_notes_backup()
    user_data_path = bundle_dir / BACKUP_BUNDLE_USER_DATA_NAME
    notes_json_path = bundle_dir / BACKUP_BUNDLE_NOTES_JSON_NAME
    notes_docx_path = bundle_dir / BACKUP_BUNDLE_NOTES_DOCX_NAME
    manifest_path = bundle_dir / BACKUP_BUNDLE_MANIFEST_NAME
    _write_json(user_data_path, user_data_backup, indent=2)
    _write_json(notes_json_path, notes_backup, indent=2)
    export_notes_docx(notes_docx_path)
    _write_json(manifest_path, _build_backup_bundle_manifest(user_data_backup, notes_backup), indent=2)
    return {
        'folder': str(bundle_dir),
        'manifest_path': str(manifest_path),
        'user_data_path': str(user_data_path),
        'notes_json_path': str(notes_json_path),
        'notes_path': str(notes_docx_path),
    }


def create_rollback_backup_bundle(*, reason: str='before_import') -> dict[str, str]:
    """Create an automatic rollback bundle under local app data."""
    safe_reason = ''.join(char if char.isalnum() or char in {'_', '-'} else '_' for char in str(reason or 'before_import')).strip('_') or 'before_import'
    return export_user_data_bundle(ROLLBACK_BUNDLES_DIR, prefix=f'rollback_{safe_reason}')


def _json_block(data: Any) -> Any:
    """Render a JSON code block for human and AI-readable exports."""
    return f"```json\n{json.dumps(data, indent=2)}\n```"


def build_ai_user_data_export() -> str:
    """Build a Markdown export optimized for human and AI analysis."""
    payload = build_user_data_backup()
    lines = [
        '# Budget Terminal User Data',
        '',
        f"- Exported at: {payload.get('exported_at', '')}",
        f"- Main portfolio: {payload.get('main_portfolio_id', DEFAULT_MAIN_PORTFOLIO_ID)}",
        f"- Active portfolio: {payload.get('active_portfolio_id', DEFAULT_MAIN_PORTFOLIO_ID)}",
        '',
        '## Portfolio Summary',
        '',
    ]
    portfolios = payload.get('portfolios', {})
    portfolio_order = _normalize_portfolio_order(payload.get('portfolio_order'), portfolios)
    for portfolio_id in portfolio_order:
        entry = portfolios.get(portfolio_id, {})
        name = str(entry.get('name', DEFAULT_PORTFOLIO_NAMES.get(portfolio_id, portfolio_id)) or DEFAULT_PORTFOLIO_NAMES.get(portfolio_id, portfolio_id))
        tickers = list(entry.get('portfolio', []))
        tracker = dict(entry.get('portfolio_tracker', {})) if isinstance(entry.get('portfolio_tracker', {}), dict) else {}
        options_data = list(entry.get('options_tracker', [])) if isinstance(entry.get('options_tracker', []), list) else []
        lines.extend([
            f"### {name} ({portfolio_id})",
            '',
            f"- Ticker count: {len(tickers)}",
            f"- Tracker rows: {len(tracker)}",
            f"- Options positions: {len(options_data)}",
            f"- Chart slots: {', '.join(list(entry.get('chart_slots', [])))}",
            '',
            '**Tickers**',
            '',
            _json_block(tickers),
            '',
            '**Tracker Data**',
            '',
            _json_block(tracker),
            '',
            '**Options Data**',
            '',
            _json_block(options_data),
            '',
        ])
    lines.extend([
        '## Net Worth',
        '',
        _json_block(payload.get('net_worth', {'cash': [], 'debt': []})),
        '',
        '## Notes',
        '',
        _json_block(payload.get('notes', [])),
        '',
        '## Charts Page Settings',
        '',
        _json_block(payload.get('chart_page', DEFAULT_CHART_PAGE_SETTINGS)),
        '',
        '## Compare Presets',
        '',
        _json_block(payload.get('compare_presets', payload.get('chart_page', {}).get('compare_presets', []))),
        '',
        '## Dashboard Chart Settings',
        '',
        _json_block(payload.get('dashboard_chart', DEFAULT_DASHBOARD_CHART_SETTINGS)),
        '',
        '## Stocks Page Settings',
        '',
        _json_block(payload.get('stocks_page', DEFAULT_STOCKS_PAGE_SETTINGS)),
        '',
        '## Portfolio Metrics Settings',
        '',
        _json_block(payload.get('portfolio_metrics', DEFAULT_PORTFOLIO_METRICS_SETTINGS)),
        '',
        '## Multi Charts Settings',
        '',
        _json_block(payload.get('multi_charts', DEFAULT_MULTI_CHARTS_SETTINGS)),
        '',
        '## Theme Settings',
        '',
        _json_block(payload.get('theme', DEFAULT_THEME_SETTINGS)),
        '',
        '## Options Chain Settings',
        '',
        _json_block(payload.get('options_chain', DEFAULT_OPTIONS_CHAIN_SETTINGS)),
        '',
        '## Full Normalized Payload',
        '',
        _json_block(payload),
        '',
    ])
    return '\n'.join(lines)


def export_ai_user_data(path: Any) -> None:
    """Write an AI-friendly Markdown export of the current user data."""
    Path(path).write_text(build_ai_user_data_export(), encoding='utf-8')


def _validate_backup_payload(payload: Any, *, preserve_existing_notes: bool=False) -> Any:
    """Validate imported backup data and return a normalized payload."""
    if not isinstance(payload, dict):
        raise ValueError('Backup file must contain a JSON object.')
    has_multi_portfolios = isinstance(payload.get('portfolios'), dict)
    has_legacy_payload = any((key in payload for key in ('portfolio', 'portfolio_tracker', 'options_tracker')))
    if not has_multi_portfolios and not has_legacy_payload:
        raise ValueError('Backup file is missing portfolio data.')
    if 'net_worth' in payload:
        networth_payload = payload.get('net_worth')
        if not isinstance(networth_payload, dict):
            raise ValueError('Backup net worth data must be a JSON object.')
        if not isinstance(networth_payload.get('cash', []), list) or not isinstance(networth_payload.get('pension_insurance', []), list) or not isinstance(networth_payload.get('debt', []), list):
            raise ValueError('Backup net worth data must include cash, pension_insurance, and debt lists.')
    existing_notes = load_notes_data() if preserve_existing_notes and 'notes' not in payload else None
    return _normalize_user_data_document(payload, existing_notes=existing_notes)


def _summarize_user_data_payload(payload: Any) -> dict[str, Any]:
    """Build a compact summary for a normalized user-data payload."""
    normalized = _validate_backup_payload(payload, preserve_existing_notes=False)
    portfolio_state = _normalize_multi_portfolio_state(normalized)
    notes = normalized.get('notes', [])
    return {
        'portfolios': len(portfolio_state.get('portfolio_order', [])),
        'notes': len(notes) if isinstance(notes, list) else 0,
        'note_images': _count_note_images(notes),
        'exported_at': str(normalized.get('exported_at', '') or payload.get('exported_at', '')) if isinstance(payload, dict) else '',
        'app_version': str(payload.get('app_version', '') or '') if isinstance(payload, dict) else '',
    }


def _summarize_notes_payload(payload: Any) -> dict[str, Any]:
    """Build a compact summary for a notes backup payload."""
    validated = _validate_notes_backup_payload(payload)
    notes = validated.get('notes', [])
    return {
        'portfolios': 0,
        'notes': len(notes) if isinstance(notes, list) else 0,
        'note_images': _count_note_images(notes),
        'exported_at': str(validated.get('exported_at', '') or validated.get('imported_at', '')),
        'app_version': '',
    }


def _extract_notes_payload_from_user_data(payload: Any) -> Any:
    """Extract a notes-only backup payload from a user-data backup when possible."""
    normalized = _validate_backup_payload(payload, preserve_existing_notes=False)
    notes = normalized.get('notes', [])
    if not isinstance(notes, list):
        return None
    return {
        'version': NOTES_BACKUP_VERSION,
        'imported_at': datetime.datetime.now(datetime.timezone.utc).isoformat(),
        'notes': notes,
    }


def inspect_import_source(path: Any) -> dict[str, Any]:
    """Inspect a user-selected backup source without mutating persisted state."""
    source_path = Path(path)
    if not source_path.exists():
        raise ValueError('Selected backup source does not exist.')

    if source_path.is_dir() or source_path.name.lower() == BACKUP_BUNDLE_MANIFEST_NAME:
        bundle_dir = source_path if source_path.is_dir() else source_path.parent
        manifest_path = bundle_dir / BACKUP_BUNDLE_MANIFEST_NAME
        manifest = _read_json(manifest_path, {}) if manifest_path.exists() else {}
        files_meta = manifest.get('files', {}) if isinstance(manifest, dict) else {}
        user_data_name = str(files_meta.get('user_data', BACKUP_BUNDLE_USER_DATA_NAME) or BACKUP_BUNDLE_USER_DATA_NAME)
        notes_name = str(files_meta.get('notes', BACKUP_BUNDLE_NOTES_JSON_NAME) or BACKUP_BUNDLE_NOTES_JSON_NAME)
        notes_docx_name = str(files_meta.get('notes_docx', BACKUP_BUNDLE_NOTES_DOCX_NAME) or BACKUP_BUNDLE_NOTES_DOCX_NAME)
        user_data_file = bundle_dir / user_data_name
        notes_file = bundle_dir / notes_name
        notes_docx_file = bundle_dir / notes_docx_name
        found_files = []
        missing_files = []
        user_data_payload = None
        notes_payload = None
        if manifest_path.exists():
            found_files.append(BACKUP_BUNDLE_MANIFEST_NAME)
        else:
            missing_files.append(BACKUP_BUNDLE_MANIFEST_NAME)
        if user_data_file.exists():
            raw_user_data_payload = _read_json(user_data_file, None)
            if raw_user_data_payload is None:
                raise ValueError(f'Unable to read backup file: {user_data_file.name}')
            user_data_payload = _validate_backup_payload(raw_user_data_payload, preserve_existing_notes=True)
            found_files.append(user_data_file.name)
        else:
            missing_files.append(user_data_file.name)
        if notes_file.exists():
            notes_payload = load_notes_backup(notes_file)
            found_files.append(notes_file.name)
        elif notes_docx_file.exists():
            notes_payload = load_notes_backup(notes_docx_file)
            found_files.append(notes_docx_file.name)
            missing_files.append(notes_file.name)
        else:
            missing_files.extend([notes_name, notes_docx_name])
        if notes_docx_file.exists() and notes_docx_file.name not in found_files:
            found_files.append(notes_docx_file.name)
        supported_scopes = []
        if user_data_payload is not None and notes_payload is not None:
            supported_scopes.append('full')
        if user_data_payload is not None:
            supported_scopes.append('user_data_only')
        if notes_payload is not None:
            supported_scopes.append('notes_only')
        if not supported_scopes:
            raise ValueError('Backup bundle is missing both user_data and notes backups.')
        counts = dict(manifest.get('counts', {})) if isinstance(manifest, dict) else {}
        user_summary = _summarize_user_data_payload(raw_user_data_payload) if user_data_payload is not None else {'portfolios': 0, 'notes': 0, 'note_images': 0, 'exported_at': '', 'app_version': ''}
        notes_summary = _summarize_notes_payload(notes_payload) if notes_payload is not None else {'notes': 0, 'note_images': 0, 'exported_at': '', 'app_version': ''}
        return {
            'source_kind': 'bundle',
            'source_path': str(source_path),
            'display_name': str(bundle_dir),
            'manifest': manifest if isinstance(manifest, dict) else {},
            'user_data_payload': user_data_payload,
            'notes_payload': notes_payload,
            'supported_scopes': supported_scopes,
            'exported_at': str((manifest.get('exported_at') if isinstance(manifest, dict) else '') or user_summary.get('exported_at', '') or notes_summary.get('exported_at', '')),
            'app_version': str((manifest.get('app_version') if isinstance(manifest, dict) else '') or user_summary.get('app_version', '')),
            'portfolio_count': int(counts.get('portfolios', user_summary.get('portfolios', 0)) or 0),
            'notes_count': int(counts.get('notes', notes_summary.get('notes', 0)) or 0),
            'image_count': int(counts.get('note_images', notes_summary.get('note_images', 0)) or 0),
            'found_files': found_files,
            'missing_files': missing_files,
        }

    if source_path.suffix.lower() == '.docx':
        notes_payload = load_notes_backup(source_path)
        notes_summary = _summarize_notes_payload(notes_payload)
        return {
            'source_kind': 'notes',
            'source_path': str(source_path),
            'display_name': str(source_path),
            'manifest': {},
            'user_data_payload': None,
            'notes_payload': notes_payload,
            'supported_scopes': ['notes_only'],
            'exported_at': str(notes_summary.get('exported_at', '')),
            'app_version': '',
            'portfolio_count': 0,
            'notes_count': int(notes_summary.get('notes', 0) or 0),
            'image_count': int(notes_summary.get('note_images', 0) or 0),
            'found_files': [source_path.name],
            'missing_files': [],
        }

    payload = _read_json(source_path, None)
    if payload is None:
        raise ValueError('Unable to read backup file.')
    if source_path.name.lower() == BACKUP_BUNDLE_MANIFEST_NAME:
        return inspect_import_source(source_path.parent)
    if isinstance(payload, dict) and isinstance(payload.get('notes'), list) and not isinstance(payload.get('portfolios'), dict) and not any((key in payload for key in ('portfolio', 'portfolio_tracker', 'options_tracker'))):
        notes_payload = _validate_notes_backup_payload(payload)
        notes_summary = _summarize_notes_payload(notes_payload)
        return {
            'source_kind': 'notes',
            'source_path': str(source_path),
            'display_name': str(source_path),
            'manifest': {},
            'user_data_payload': None,
            'notes_payload': notes_payload,
            'supported_scopes': ['notes_only'],
            'exported_at': str(notes_summary.get('exported_at', '')),
            'app_version': '',
            'portfolio_count': 0,
            'notes_count': int(notes_summary.get('notes', 0) or 0),
            'image_count': int(notes_summary.get('note_images', 0) or 0),
            'found_files': [source_path.name],
            'missing_files': [],
        }
    user_data_payload = _validate_backup_payload(payload, preserve_existing_notes=True)
    user_summary = _summarize_user_data_payload(payload)
    supported_scopes = ['full', 'user_data_only']
    extracted_notes_payload = None
    try:
        extracted_notes_payload = _extract_notes_payload_from_user_data(payload)
    except Exception:
        extracted_notes_payload = None
    if extracted_notes_payload is not None and _summarize_notes_payload(extracted_notes_payload).get('notes', 0):
        supported_scopes.append('notes_only')
    return {
        'source_kind': 'user_data',
        'source_path': str(source_path),
        'display_name': str(source_path),
        'manifest': {},
        'user_data_payload': user_data_payload,
        'notes_payload': extracted_notes_payload,
        'supported_scopes': supported_scopes,
        'exported_at': str(user_summary.get('exported_at', '')),
        'app_version': str(user_summary.get('app_version', '')),
        'portfolio_count': int(user_summary.get('portfolios', 0) or 0),
        'notes_count': int(user_summary.get('notes', 0) or 0),
        'image_count': int(user_summary.get('note_images', 0) or 0),
        'found_files': [source_path.name],
        'missing_files': [],
    }


def load_user_data_backup(path: Any) -> Any:
    """Load and validate a single-file backup payload."""
    payload = _read_json(path, None)
    if payload is None:
        raise ValueError('Unable to read backup file.')
    return _validate_backup_payload(payload, preserve_existing_notes=True)


def apply_user_data_backup(payload: Any) -> Any:
    """Persist validated backup data and return the normalized state."""
    normalized = _validate_backup_payload(payload, preserve_existing_notes=True)
    return _save_user_data_document(normalized)


def apply_import_source(import_source: Any, *, scope: str='full') -> Any:
    """Apply an inspected import source and return the current normalized user-data state."""
    source = import_source if isinstance(import_source, dict) else {}
    user_data_payload = source.get('user_data_payload')
    notes_payload = source.get('notes_payload')
    supported_scopes = list(source.get('supported_scopes', [])) if isinstance(source.get('supported_scopes', []), list) else []
    if scope not in supported_scopes:
        raise ValueError(f'Import scope "{scope}" is not available for this backup source.')
    if scope == 'full':
        if user_data_payload is None:
            raise ValueError('Full restore requires user data.')
        apply_user_data_backup(user_data_payload)
        if notes_payload is not None:
            apply_notes_backup(notes_payload)
        return _load_user_data_document()
    if scope == 'user_data_only':
        if user_data_payload is None:
            raise ValueError('This backup source does not include user data.')
        return apply_user_data_backup(user_data_payload)
    if scope == 'notes_only':
        if notes_payload is None:
            raise ValueError('This backup source does not include notes.')
        apply_notes_backup(notes_payload)
        return _load_user_data_document()
    raise ValueError(f'Unknown import scope: {scope}')


def reset_user_data(chart_slots: Any=None) -> Any:
    """Persist a cleared user-data state while preserving chart slots."""
    normalized = {
        'main_portfolio_id': DEFAULT_MAIN_PORTFOLIO_ID,
        'active_portfolio_id': DEFAULT_MAIN_PORTFOLIO_ID,
        'portfolio_order': [DEFAULT_MAIN_PORTFOLIO_ID],
        'portfolios': {
            DEFAULT_MAIN_PORTFOLIO_ID: _default_portfolio_entry(DEFAULT_MAIN_PORTFOLIO_ID, chart_slots)
        },
        'fundamentals_page': _normalize_fundamentals_page_settings(DEFAULT_FUNDAMENTALS_PAGE_SETTINGS),
        'chart_page': DEFAULT_CHART_PAGE_SETTINGS.copy(),
        'dashboard_chart': DEFAULT_DASHBOARD_CHART_SETTINGS.copy(),
        'stocks_page': DEFAULT_STOCKS_PAGE_SETTINGS.copy(),
        'portfolio_metrics': DEFAULT_PORTFOLIO_METRICS_SETTINGS.copy(),
        'multi_charts': DEFAULT_MULTI_CHARTS_SETTINGS.copy(),
        'youtube': DEFAULT_YOUTUBE_SETTINGS.copy(),
        'net_worth': {'cash': [], 'pension_insurance': [], 'debt': []},
        'notes': list(DEFAULT_NOTES),
        'theme': DEFAULT_THEME_SETTINGS.copy(),
        'options_chain': DEFAULT_OPTIONS_CHAIN_SETTINGS.copy(),
        'time_12h': False,
    }
    return apply_user_data_backup(normalized)


def load_app_config() -> Any:
    """Load app-level config data."""
    document = _load_user_data_document()
    return {
        'portfolio_state': {
            'main_portfolio_id': document['main_portfolio_id'],
            'active_portfolio_id': document.get('active_portfolio_id', document['main_portfolio_id']),
            'portfolio_order': list(document.get('portfolio_order', [document['main_portfolio_id']])),
            'portfolios': {portfolio_id: {'name': entry.get('name')} for portfolio_id, entry in document.get('portfolios', {}).items()},
        },
        'theme': dict(document.get('theme', DEFAULT_THEME_SETTINGS)),
        'fundamentals_page': dict(document.get('fundamentals_page', DEFAULT_FUNDAMENTALS_PAGE_SETTINGS)),
        'chart_page': dict(document.get('chart_page', DEFAULT_CHART_PAGE_SETTINGS)),
        'dashboard_chart': dict(document.get('dashboard_chart', DEFAULT_DASHBOARD_CHART_SETTINGS)),
        'stocks_page': dict(document.get('stocks_page', DEFAULT_STOCKS_PAGE_SETTINGS)),
        'portfolio_metrics': dict(document.get('portfolio_metrics', DEFAULT_PORTFOLIO_METRICS_SETTINGS)),
        'multi_charts': dict(document.get('multi_charts', DEFAULT_MULTI_CHARTS_SETTINGS)),
        'youtube': dict(document.get('youtube', DEFAULT_YOUTUBE_SETTINGS)),
        'options_chain': dict(document.get('options_chain', DEFAULT_OPTIONS_CHAIN_SETTINGS)),
        'time_12h': bool(document.get('time_12h', False)),
    }


def save_app_config(data: Any) -> None:
    """Persist app-level config data."""
    current = _load_user_data_document()
    saved = data if isinstance(data, dict) else {}
    portfolio_state = saved.get('portfolio_state', {})
    if isinstance(portfolio_state, dict):
        order = _normalize_portfolio_order(
            portfolio_state.get('portfolio_order'),
            portfolio_state.get('portfolios', current.get('portfolios', {})),
        )
        current['portfolio_order'] = list(order)
        current['main_portfolio_id'] = _normalize_selected_portfolio_id(portfolio_state.get('main_portfolio_id', current['main_portfolio_id']), order, order[0])
        current['active_portfolio_id'] = _normalize_selected_portfolio_id(
            portfolio_state.get('active_portfolio_id', current.get('active_portfolio_id', current['main_portfolio_id'])),
            order,
            current['main_portfolio_id'],
        )
        names = _normalize_portfolio_catalog(portfolio_state.get('portfolios'), order)
        portfolio_map = current.get('portfolios', {})
        current['portfolios'] = {}
        for portfolio_id in order:
            entry = portfolio_map.get(portfolio_id, _default_portfolio_entry(portfolio_id))
            if not isinstance(entry, dict):
                entry = _default_portfolio_entry(portfolio_id)
            entry['name'] = names.get(portfolio_id, {}).get('name', entry.get('name'))
            current['portfolios'][portfolio_id] = entry
    if 'theme' in saved:
        current['theme'] = _normalize_theme_payload(saved.get('theme'))
    if 'fundamentals_page' in saved:
        current['fundamentals_page'] = _normalize_fundamentals_page_settings(saved.get('fundamentals_page'))
    if 'chart_page' in saved:
        current['chart_page'] = _normalize_chart_page_settings(saved.get('chart_page'))
    if 'dashboard_chart' in saved:
        current['dashboard_chart'] = _normalize_dashboard_chart_settings(saved.get('dashboard_chart'))
    if 'stocks_page' in saved:
        current['stocks_page'] = _normalize_stocks_page_settings(saved.get('stocks_page'))
    if 'portfolio_metrics' in saved:
        current['portfolio_metrics'] = _normalize_portfolio_metrics_settings(saved.get('portfolio_metrics'))
    if 'multi_charts' in saved:
        current['multi_charts'] = _normalize_multi_charts_settings(saved.get('multi_charts'))
    if 'youtube' in saved:
        current['youtube'] = _normalize_youtube_settings(saved.get('youtube'))
    if 'options_chain' in saved:
        current['options_chain'] = _normalize_options_chain_payload(saved.get('options_chain'))
    if 'time_12h' in saved:
        current['time_12h'] = bool(saved['time_12h'])
    _save_user_data_document(current)


def _normalize_indicator_list(raw_indicators: Any, default_indicators: Any, *, ensure_ma200: bool=False) -> Any:
    """Normalize indicator selections into the supported canonical labels."""
    indicators_source = raw_indicators if isinstance(raw_indicators, list) else list(default_indicators)
    indicators = []
    for name in indicators_source:
        text = str(name or '').strip()
        normalized = text.upper().replace(' ', '')
        if normalized == 'VOLUME':
            text = 'Volume'
        elif normalized == 'RSI':
            text = 'RSI'
        elif normalized in ('200MA', 'MA200'):
            text = '200 MA'
        else:
            text = ''
        if text and text not in indicators:
            indicators.append(text)
    if ensure_ma200 and '200 MA' not in indicators:
        indicators.append('200 MA')
    ordered = []
    for indicator in ('Volume', 'RSI', '200 MA'):
        if indicator in indicators and indicator not in ordered:
            ordered.append(indicator)
    return ordered or list(default_indicators)


def _normalize_chart_page_settings(settings: Any) -> Any:
    """Normalize persisted state for the dedicated Charts page."""
    saved = settings if isinstance(settings, dict) else {}
    symbol = str(saved.get('symbol', DEFAULT_CHART_PAGE_SETTINGS['symbol']) or DEFAULT_CHART_PAGE_SETTINGS['symbol']).upper()
    timeframe_label = str(saved.get('timeframe_label', DEFAULT_CHART_PAGE_SETTINGS['timeframe_label']) or DEFAULT_CHART_PAGE_SETTINGS['timeframe_label'])
    compare_interval_label = str(saved.get('compare_interval_label', DEFAULT_CHART_PAGE_SETTINGS['compare_interval_label']) or DEFAULT_CHART_PAGE_SETTINGS['compare_interval_label']).strip()
    if compare_interval_label not in {'1 Day', '1 Week'}:
        compare_interval_label = DEFAULT_CHART_PAGE_SETTINGS['compare_interval_label']
    compare_range_label = str(saved.get('compare_range_label', DEFAULT_CHART_PAGE_SETTINGS['compare_range_label']) or DEFAULT_CHART_PAGE_SETTINGS['compare_range_label']).strip().upper()
    if compare_range_label not in {'5Y', '3Y', '1Y', 'YTD', '3M', '1M'}:
        compare_range_label = DEFAULT_CHART_PAGE_SETTINGS['compare_range_label']
    raw_watchlist = saved.get('watchlist', [])
    if not isinstance(raw_watchlist, list):
        raw_watchlist = []
    watchlist = _normalize_unique_symbol_list(raw_watchlist)
    raw_compare_symbols = saved.get('compare_symbols', [])
    if not isinstance(raw_compare_symbols, list):
        raw_compare_symbols = []
    compare_symbols = _normalize_unique_symbol_list(raw_compare_symbols)
    compare_presets = _normalize_compare_presets(saved.get('compare_presets', []))
    raw_multi_interval_labels = saved.get('multi_interval_labels', [])
    if not isinstance(raw_multi_interval_labels, list):
        raw_multi_interval_labels = []
    multi_interval_labels = []
    valid_multi_interval_labels = {'1 Minute', '5 Minutes', '15 Minutes', '1 Hour', '1 Day', '1 Week', '1 Month'}
    for value in raw_multi_interval_labels:
        label = str(value or '').strip()
        if label in valid_multi_interval_labels and label not in multi_interval_labels:
            multi_interval_labels.append(label)
    indicators = _normalize_indicator_list(saved.get('indicators'), DEFAULT_CHART_PAGE_SETTINGS['indicators'])
    auto_value = saved.get('auto', DEFAULT_CHART_PAGE_SETTINGS['auto'])
    auto_enabled = bool(auto_value) if isinstance(auto_value, bool | int) else DEFAULT_CHART_PAGE_SETTINGS['auto']
    return {
        'symbol': symbol,
        'timeframe_label': timeframe_label,
        'compare_interval_label': compare_interval_label,
        'compare_range_label': compare_range_label,
        'watchlist': watchlist,
        'compare_symbols': compare_symbols,
        'compare_presets': compare_presets,
        'multi_interval_labels': multi_interval_labels,
        'indicators': indicators,
        'auto': auto_enabled,
    }


def _normalize_fundamentals_selection_rows(values: Any) -> list[str]:
    """Normalize one ordered Fundamentals row-selection list."""
    rows = []
    if not isinstance(values, list):
        return rows
    for value in values:
        row = str(value or '').strip()
        if row and row not in rows:
            rows.append(row)
    return rows


def _normalize_fundamentals_custom_selections(values: Any, *, last_ticker: str='', legacy_custom_panels: Any=None) -> dict[str, dict[str, list[str]]]:
    """Normalize per-ticker Fundamentals checklist selections, with migration support."""
    normalized = {}
    raw = values if isinstance(values, dict) else {}
    for ticker_key, selection in raw.items():
        ticker = str(ticker_key or '').upper().strip()
        if not ticker or not isinstance(selection, dict):
            continue
        family_map = {
            'financials': _normalize_fundamentals_selection_rows(selection.get('financials', [])),
            'cashflow': _normalize_fundamentals_selection_rows(selection.get('cashflow', [])),
            'balance_sheet': _normalize_fundamentals_selection_rows(selection.get('balance_sheet', [])),
        }
        if any(family_map.values()):
            normalized[ticker] = family_map
    if normalized:
        return normalized
    ticker = str(last_ticker or '').upper().strip()
    if not ticker or not isinstance(legacy_custom_panels, list):
        return {}
    migrated = {
        'financials': [],
        'cashflow': [],
        'balance_sheet': [],
    }
    for entry in legacy_custom_panels:
        panel = entry if isinstance(entry, dict) else {}
        if str(panel.get('source', '') or '').strip().lower() != 'statement_row':
            continue
        family = str(panel.get('statement_family', 'financials') or 'financials').strip().lower()
        if family not in migrated:
            continue
        row = str(panel.get('statement_row', '') or '').strip()
        if row and row not in migrated[family]:
            migrated[family].append(row)
    return {ticker: migrated} if any(migrated.values()) else {}


def _normalize_fundamentals_page_settings(settings: Any) -> dict[str, Any]:
    """Normalize persisted state for the Fundamentals page."""
    saved = settings if isinstance(settings, dict) else {}
    last_ticker = str(saved.get('last_ticker', DEFAULT_FUNDAMENTALS_PAGE_SETTINGS['last_ticker']) or '').upper().strip()
    selected_configuration = str(
        saved.get('selected_configuration', DEFAULT_FUNDAMENTALS_PAGE_SETTINGS['selected_configuration'])
        or DEFAULT_FUNDAMENTALS_PAGE_SETTINGS['selected_configuration']
    ).strip().lower()
    if selected_configuration not in {'default', 'custom'}:
        selected_configuration = DEFAULT_FUNDAMENTALS_PAGE_SETTINGS['selected_configuration']
    return {
        'last_ticker': last_ticker,
        'selected_configuration': selected_configuration,
        'custom_selections_by_ticker': _normalize_fundamentals_custom_selections(
            saved.get('custom_selections_by_ticker', DEFAULT_FUNDAMENTALS_PAGE_SETTINGS['custom_selections_by_ticker']),
            last_ticker=last_ticker,
            legacy_custom_panels=saved.get('custom_panels'),
        ),
    }


def _normalize_stocks_page_settings(settings: Any) -> Any:
    """Normalize persisted state for the Stocks page."""
    saved = settings if isinstance(settings, dict) else {}
    symbol = str(saved.get('symbol', DEFAULT_STOCKS_PAGE_SETTINGS['symbol']) or DEFAULT_STOCKS_PAGE_SETTINGS['symbol']).upper().strip()
    auto_value = saved.get('auto', DEFAULT_STOCKS_PAGE_SETTINGS['auto'])
    auto_enabled = bool(auto_value) if isinstance(auto_value, bool | int) else DEFAULT_STOCKS_PAGE_SETTINGS['auto']
    mfi_value = saved.get('mfi_enabled', DEFAULT_STOCKS_PAGE_SETTINGS['mfi_enabled'])
    mfi_enabled = bool(mfi_value) if isinstance(mfi_value, bool | int) else DEFAULT_STOCKS_PAGE_SETTINGS['mfi_enabled']
    raw_main_splitter = saved.get('main_splitter_sizes', DEFAULT_STOCKS_PAGE_SETTINGS['main_splitter_sizes'])
    main_splitter_sizes = []
    if isinstance(raw_main_splitter, list):
        for value in raw_main_splitter[:3]:
            try:
                size = max(int(value), 1)
            except (TypeError, ValueError):
                size = 0
            if size > 0:
                main_splitter_sizes.append(size)
    if len(main_splitter_sizes) != 3:
        main_splitter_sizes = list(DEFAULT_STOCKS_PAGE_SETTINGS['main_splitter_sizes'])

    raw_left_splitter = saved.get('left_splitter_sizes', DEFAULT_STOCKS_PAGE_SETTINGS['left_splitter_sizes'])
    left_splitter_sizes = []
    if isinstance(raw_left_splitter, list):
        for value in raw_left_splitter[:3]:
            try:
                size = max(int(value), 1)
            except (TypeError, ValueError):
                size = 0
            if size > 0:
                left_splitter_sizes.append(size)
    if len(left_splitter_sizes) != 3:
        left_splitter_sizes = list(DEFAULT_STOCKS_PAGE_SETTINGS['left_splitter_sizes'])

    raw_middle_splitter = saved.get('middle_splitter_sizes', DEFAULT_STOCKS_PAGE_SETTINGS['middle_splitter_sizes'])
    middle_splitter_sizes = []
    if isinstance(raw_middle_splitter, list):
        for value in raw_middle_splitter[:2]:
            try:
                size = max(int(value), 1)
            except (TypeError, ValueError):
                size = 0
            if size > 0:
                middle_splitter_sizes.append(size)
    if len(middle_splitter_sizes) != 2:
        middle_splitter_sizes = list(DEFAULT_STOCKS_PAGE_SETTINGS['middle_splitter_sizes'])

    return {
        'symbol': symbol or DEFAULT_STOCKS_PAGE_SETTINGS['symbol'],
        'auto': auto_enabled,
        'mfi_enabled': mfi_enabled,
        'main_splitter_sizes': main_splitter_sizes,
        'left_splitter_sizes': left_splitter_sizes,
        'middle_splitter_sizes': middle_splitter_sizes,
    }


def _normalize_portfolio_metrics_settings(settings: Any) -> Any:
    """Normalize persisted state for the Portfolio Metrics sub-tab."""
    saved = settings if isinstance(settings, dict) else {}
    benchmark_symbol = str(
        saved.get('benchmark_symbol', DEFAULT_PORTFOLIO_METRICS_SETTINGS['benchmark_symbol'])
        or DEFAULT_PORTFOLIO_METRICS_SETTINGS['benchmark_symbol']
    ).upper().strip()
    lookback_key = str(
        saved.get('lookback_key', DEFAULT_PORTFOLIO_METRICS_SETTINGS['lookback_key'])
        or DEFAULT_PORTFOLIO_METRICS_SETTINGS['lookback_key']
    ).strip().lower()
    if lookback_key not in PORTFOLIO_METRICS_LOOKBACK_CHOICES:
        lookback_key = DEFAULT_PORTFOLIO_METRICS_SETTINGS['lookback_key']
    return {
        'benchmark_symbol': benchmark_symbol or DEFAULT_PORTFOLIO_METRICS_SETTINGS['benchmark_symbol'],
        'lookback_key': lookback_key,
    }


def _normalize_multi_charts_settings(settings: Any) -> Any:
    """Normalize persisted state for the Multi Charts page."""
    saved = settings if isinstance(settings, dict) else {}
    return {
        'custom_symbols': _normalize_unique_symbol_list(saved.get('custom_symbols', DEFAULT_MULTI_CHARTS_SETTINGS['custom_symbols'])),
        'order': _normalize_unique_symbol_list(saved.get('order', DEFAULT_MULTI_CHARTS_SETTINGS['order'])),
    }


def _normalize_youtube_settings(settings: Any) -> Any:
    """Normalize persisted state for the YouTube page."""
    saved = settings if isinstance(settings, dict) else {}
    try:
        sort_column = int(saved.get('sort_column', DEFAULT_YOUTUBE_SETTINGS['sort_column']))
    except (TypeError, ValueError):
        sort_column = DEFAULT_YOUTUBE_SETTINGS['sort_column']
    if sort_column < -1 or sort_column > 6:
        sort_column = DEFAULT_YOUTUBE_SETTINGS['sort_column']
    descending_value = saved.get('sort_descending', DEFAULT_YOUTUBE_SETTINGS['sort_descending'])
    sort_descending = bool(descending_value) if isinstance(descending_value, bool | int) else DEFAULT_YOUTUBE_SETTINGS['sort_descending']
    return {
        'sort_column': sort_column,
        'sort_descending': sort_descending,
    }


def _normalize_dashboard_chart_settings(settings: Any) -> Any:
    """Normalize persisted state for the dashboard chart workstation."""
    saved = settings if isinstance(settings, dict) else {}
    symbol = str(saved.get('symbol', DEFAULT_DASHBOARD_CHART_SETTINGS['symbol']) or DEFAULT_DASHBOARD_CHART_SETTINGS['symbol']).upper().strip()
    timeframe_label = str(saved.get('timeframe_label', DEFAULT_DASHBOARD_CHART_SETTINGS['timeframe_label']) or DEFAULT_DASHBOARD_CHART_SETTINGS['timeframe_label']).strip()
    indicators = _normalize_indicator_list(
        saved.get('indicators'),
        DEFAULT_DASHBOARD_CHART_SETTINGS['indicators'],
        ensure_ma200=True,
    )
    auto_value = saved.get('auto', DEFAULT_DASHBOARD_CHART_SETTINGS['auto'])
    auto_enabled = bool(auto_value) if isinstance(auto_value, bool | int) else DEFAULT_DASHBOARD_CHART_SETTINGS['auto']
    raw_splitter_sizes = saved.get('splitter_sizes', DEFAULT_DASHBOARD_CHART_SETTINGS['splitter_sizes'])
    splitter_sizes = []
    if isinstance(raw_splitter_sizes, list):
        for value in raw_splitter_sizes[:2]:
            try:
                size = max(int(value), 1)
            except (TypeError, ValueError):
                size = 0
            if size > 0:
                splitter_sizes.append(size)
    if len(splitter_sizes) != 2:
        splitter_sizes = list(DEFAULT_DASHBOARD_CHART_SETTINGS['splitter_sizes'])
    raw_main_splitter = saved.get('main_splitter_sizes', DEFAULT_DASHBOARD_CHART_SETTINGS['main_splitter_sizes'])
    main_splitter_sizes = []
    if isinstance(raw_main_splitter, list):
        for value in raw_main_splitter[:2]:
            try:
                size = max(int(value), 1)
            except (TypeError, ValueError):
                size = 0
            if size > 0:
                main_splitter_sizes.append(size)
    if len(main_splitter_sizes) != 2:
        main_splitter_sizes = list(DEFAULT_DASHBOARD_CHART_SETTINGS['main_splitter_sizes'])
    return {
        'symbol': symbol or DEFAULT_DASHBOARD_CHART_SETTINGS['symbol'],
        'timeframe_label': timeframe_label or DEFAULT_DASHBOARD_CHART_SETTINGS['timeframe_label'],
        'indicators': indicators,
        'auto': auto_enabled,
        'splitter_sizes': splitter_sizes,
        'main_splitter_sizes': main_splitter_sizes,
    }


def normalize_dashboard_chart_settings(settings: Any) -> Any:
    """Public wrapper for dashboard chart state normalization."""
    return _normalize_dashboard_chart_settings(settings)


def load_portfolio_preferences() -> Any:
    """Load portfolio metadata and selected main/active portfolio ids from app config."""
    config = load_app_config()
    saved = config.get('portfolio_state', {})
    if not isinstance(saved, dict):
        saved = {}
    order = _normalize_portfolio_order(saved.get('portfolio_order'), saved.get('portfolios'))
    catalog = _normalize_portfolio_catalog(saved.get('portfolios'), order)
    main_portfolio_id = _normalize_selected_portfolio_id(saved.get('main_portfolio_id'), order, order[0])
    active_portfolio_id = _normalize_selected_portfolio_id(saved.get('active_portfolio_id', main_portfolio_id), order, main_portfolio_id)
    return {'main_portfolio_id': main_portfolio_id, 'active_portfolio_id': active_portfolio_id, 'portfolio_order': order, 'portfolios': catalog}


def save_portfolio_preferences(settings: Any) -> Any:
    """Persist portfolio metadata and selected main/active portfolio ids in app config."""
    current = load_app_config()
    state = load_portfolio_preferences()
    if isinstance(settings, dict):
        order = _normalize_portfolio_order(settings.get('portfolio_order', state.get('portfolio_order')), settings.get('portfolios', state['portfolios']))
        state['portfolio_order'] = order
        state['main_portfolio_id'] = _normalize_selected_portfolio_id(settings.get('main_portfolio_id', state['main_portfolio_id']), order, order[0])
        state['active_portfolio_id'] = _normalize_selected_portfolio_id(settings.get('active_portfolio_id', state.get('active_portfolio_id', state['main_portfolio_id'])), order, state['main_portfolio_id'])
        state['portfolios'] = _normalize_portfolio_catalog(settings.get('portfolios', state['portfolios']), order)
    current['portfolio_state'] = {
        'main_portfolio_id': state['main_portfolio_id'],
        'active_portfolio_id': state['active_portfolio_id'],
        'portfolio_order': list(state.get('portfolio_order', [state['main_portfolio_id']])),
        'portfolios': state['portfolios'],
    }
    save_app_config(current)
    return state


def load_theme_settings() -> Any:
    """Load persisted UI theme preferences."""
    config = load_app_config()
    saved = config.get('theme', {})
    if not isinstance(saved, dict):
        saved = {}
    return {'selected_theme': _normalize_theme_setting(saved.get('selected_theme'))}


def save_theme_settings(settings: Any) -> Any:
    """Persist UI theme preferences in app config."""
    current = load_app_config()
    payload = load_theme_settings()
    if isinstance(settings, dict):
        payload['selected_theme'] = _normalize_theme_setting(settings.get('selected_theme', payload['selected_theme']))
    current['theme'] = payload
    save_app_config(current)
    return payload


def load_chart_page_settings() -> Any:
    """Load persisted state for the dedicated Charts page."""
    config = load_app_config()
    return _normalize_chart_page_settings(config.get('chart_page', {}))


def save_chart_page_settings(settings: Any) -> Any:
    """Persist state for the dedicated Charts page."""
    current = load_app_config()
    state = _normalize_chart_page_settings(settings)
    current['chart_page'] = state
    save_app_config(current)
    return state


def load_fundamentals_page_settings() -> Any:
    """Load persisted state for the Fundamentals page."""
    config = load_app_config()
    return _normalize_fundamentals_page_settings(config.get('fundamentals_page', {}))


def save_fundamentals_page_settings(settings: Any) -> Any:
    """Persist state for the Fundamentals page."""
    current = load_app_config()
    state = _normalize_fundamentals_page_settings(settings)
    current['fundamentals_page'] = state
    save_app_config(current)
    return state


def load_dashboard_chart_settings() -> Any:
    """Load persisted state for the dashboard chart workstation."""
    config = load_app_config()
    return _normalize_dashboard_chart_settings(config.get('dashboard_chart', {}))


def save_dashboard_chart_settings(settings: Any) -> Any:
    """Persist state for the dashboard chart workstation."""
    current = load_app_config()
    state = _normalize_dashboard_chart_settings(settings)
    current['dashboard_chart'] = state
    save_app_config(current)
    return state


def load_stocks_page_settings() -> Any:
    """Load persisted state for the Stocks page."""
    config = load_app_config()
    return _normalize_stocks_page_settings(config.get('stocks_page', {}))


def save_stocks_page_settings(settings: Any) -> Any:
    """Persist state for the Stocks page."""
    current = load_app_config()
    state = _normalize_stocks_page_settings(settings)
    current['stocks_page'] = state
    save_app_config(current)
    return state


def load_portfolio_metrics_settings() -> Any:
    """Load persisted state for the Portfolio Metrics sub-tab."""
    config = load_app_config()
    return _normalize_portfolio_metrics_settings(config.get('portfolio_metrics', {}))


def save_portfolio_metrics_settings(settings: Any) -> Any:
    """Persist state for the Portfolio Metrics sub-tab."""
    current = load_app_config()
    state = _normalize_portfolio_metrics_settings(settings)
    current['portfolio_metrics'] = state
    save_app_config(current)
    return state


def load_multi_charts_settings() -> Any:
    """Load persisted state for the Multi Charts page."""
    config = load_app_config()
    return _normalize_multi_charts_settings(config.get('multi_charts', {}))


def save_multi_charts_settings(settings: Any) -> None:
    """Persist state for the Multi Charts page."""
    current = load_app_config()
    current['multi_charts'] = _normalize_multi_charts_settings(settings)
    save_app_config(current)


def load_youtube_settings() -> Any:
    """Load persisted state for the YouTube page."""
    config = load_app_config()
    return _normalize_youtube_settings(config.get('youtube', {}))


def save_youtube_settings(settings: Any) -> Any:
    """Persist state for the YouTube page."""
    current = load_app_config()
    state = _normalize_youtube_settings(settings)
    current['youtube'] = state
    save_app_config(current)
    return state


def load_options_chain_settings() -> Any:
    """Load persisted defaults for options-chain analytics."""
    config = load_app_config()
    saved = config.get('options_chain', {})
    if not isinstance(saved, dict):
        saved = {}
    rate = saved.get('default_risk_free_rate', DEFAULT_OPTIONS_CHAIN_SETTINGS['default_risk_free_rate'])
    try:
        rate_value = float(rate)
    except (TypeError, ValueError):
        rate_value = DEFAULT_OPTIONS_CHAIN_SETTINGS['default_risk_free_rate']
    rate_value = min(max(rate_value, 0.0), 1.0)
    return {'default_risk_free_rate': rate_value}


def save_options_chain_settings(settings: Any) -> Any:
    """Persist defaults for options-chain analytics."""
    current = load_app_config()
    state = DEFAULT_OPTIONS_CHAIN_SETTINGS.copy()
    if isinstance(settings, dict):
        raw_rate = settings.get('default_risk_free_rate', state['default_risk_free_rate'])
        try:
            state['default_risk_free_rate'] = min(max(float(raw_rate), 0.0), 1.0)
        except (TypeError, ValueError):
            state['default_risk_free_rate'] = DEFAULT_OPTIONS_CHAIN_SETTINGS['default_risk_free_rate']
    current['options_chain'] = state
    save_app_config(current)
    return state


def load_time_format() -> bool:
    """Load persisted 12h/24h time format preference."""
    config = load_app_config()
    return bool(config.get('time_12h', False))


def save_time_format(use_12h: bool) -> None:
    """Persist 12h/24h time format preference."""
    save_app_config({'time_12h': bool(use_12h)})
